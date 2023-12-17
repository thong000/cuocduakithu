import pygame
from docx import Document
from docx.shared import Inches
from datetime import datetime
import os


'''def save_screenshot(image_path, docx_file):
    with open('user.txt', 'r') as file:
        global user
        user = file.read()
    # Lấy thời gian hiện tại
    current_date = datetime.now().date()
    formatted_date = current_date.strftime("%d/%m/%Y")
    current_time = datetime.now()
    hour = current_time.hour
    minute = current_time.minute

    doc = Document(docx_file)
    time = str(user) + str(hour) + ":" + str(minute) + " ngày " + str(formatted_date)

    # Chèn ảnh và thời gian
    doc.add_paragraph(time)
    doc.add_picture(image_path, width=Inches(6.0))  # 6.0 là kích thước
    doc.save(docx_file)'''


def capture_screen(screen, user):
    screenshot = pygame.Surface((1344, 756))
    screenshot.blit(screen, (0, 0))
    pygame.image.save(screenshot, "temp.png")

    directory_path = "temp_screenshot.docx"
    docx_file = 'account/' + str(user) + '/screenshot.docx'

    # Cấp quyền truy cập cho thư mục
    try:
        # Set quyền truy cập cho owner (người sở hữu)
        os.chmod(directory_path, 0o777)  # Ví dụ: 0o700 là quyền truy cập đầy đủ cho owner

        # Set quyền truy cập cho group (nhóm)
        # os.chmod(directory_path, 0o770)  # Ví dụ: 0o770 là quyền truy cập đầy đủ cho owner và group

        # Set quyền truy cập cho others (người khác)
        # os.chmod(directory_path, 0o777)  # Ví dụ: 0o777 là quyền truy cập đầy đủ cho tất cả mọi người
    except OSError as e:
        pass

    current_date = datetime.now().date()
    formatted_date = current_date.strftime("%d/%m/%Y")
    current_time = datetime.now()
    hour = current_time.hour
    minute = current_time.minute

    doc = Document()
    time = str(user) + str(hour) + ":" + str(minute) + " ngày " + str(formatted_date)

    # Chèn ảnh và thời gian
    doc.add_paragraph(time)
    doc.add_picture("temp.png", width=Inches(6.0))  # 6.0 là kích thước
    doc.save(docx_file)
