import textwrap
import pygame
import random
import subprocess
import sys
import time
import Fruit_ninja
import PlappyBird
import Snake_game_class
import store_func
import Rule
import podium
import screenshot
import re

import function_faceid

pygame.init()
pygame.mixer.init()

write_history_active = True
with open('user.txt', 'r') as file:
    user = file.read()

# Mở file lan để xác định ngôn ngữ
with open('lan.txt', 'r') as file:
    l = file.read()
    if l == "1":
        lan = True
    if l == "0":
        lan = False

# Mở file log dể biết tình trạng đăng nhập
with open('log.txt', 'r') as file:
    content = file.read()
    if content == "1":
        lock = True
    if content == "0":
        lock = False

# Mở file coin
with open("account/" + str(user) + '/coin.txt', 'r') as c:
    global coin
    coin = int(c.read())


def read_file(file_path):
    with open(file_path, 'r') as file:
        text = file.read()
        numbers = re.findall(r'\d+', text)
        numeric_values = [int(number) for number in numbers]
        return numeric_values[0]


def read_file2(file_path):
    with open(file_path, 'r') as file:
        lines = file.readlines()
    return lines


def draw_text_list(screen, text_list, font, scroll_y):
    bg_history = pygame.image.load("background/menugamemap.png")
    bg_history = pygame.transform.scale(bg_history, (1344, 756))
    screen.blit(bg_history, (0, 0))
    column_width = 200  # Độ rộng của mỗi cột

    # Can giữ và vẽ từng dòng trong danh sách
    for i, line in enumerate(text_list):
        columns = line.strip().split(',')  # Phân chia dòng thành các cột
        for j, column in enumerate(columns):
            wrapped_lines = textwrap.wrap(column, width=22)  # Can giữ với chiều rộng tối đa 15 ký tự
            for k, wrapped_line in enumerate(wrapped_lines):
                text_surface = font.render(wrapped_line, True, (255, 255, 255))  # Màu chữ đen
                x_position = j * column_width + 10  # Vị trí của mỗi cột
                y_position = i * 20 + k * 20 - scroll_y  # Vị trí của mỗi dòng con trong cột, điều chỉnh theo lăn chuột
                screen.blit(text_surface, (x_position, y_position))

    '''pygame.display.flip()'''


# Thông số cơ bản
SCREEN_WIDTH = 1344
SCREEN_HEIGHT = 756

set_but_dis = SCREEN_WIDTH // 10
set_but_x = SCREEN_WIDTH // 10
set_but_y = SCREEN_HEIGHT // 5

exit_but_x = SCREEN_WIDTH // 40
exit_but_y = SCREEN_HEIGHT // 25
exit_w = SCREEN_WIDTH // 20
exit_h = SCREEN_HEIGHT // 10

choose_map_x = SCREEN_WIDTH // 2.5
choose_map_y = SCREEN_HEIGHT // 50

demomap_x = SCREEN_WIDTH * 2 // 3.2
demomap_y = SCREEN_HEIGHT // 2
demomap_w = SCREEN_WIDTH // 1.5
demomap_h = SCREEN_HEIGHT // 1.5

finish = SCREEN_WIDTH // 1.2

random_x1 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
random_x2 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
random_x3 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
random_x4 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
random_x5 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)

map = 0
set = 0
start_t = 0

log1st = False

channel = pygame.mixer.Channel(0)
# biến ngôn ngữ
lan = True

# Khởi tạo các biến logic để biết màn hình nào đang chạy (một lúc có 1)
set_ = True
wait = False
set1 = False
set2 = False
set3 = False
setting_bool = False
lang = False
start_bool = False
minigame_bool = False
rule_bool = False
choosenv_bool1 = False
choosenv_bool2 = False
setnv11 = False
setnv21 = False
setnv12 = False
setnv22 = False
flap = False
chonnv1 = False
chonnv2 = False
choosenv_bool3 = False
ketqua = False
history_bool = False

# Thiết lập màn hình
screen = pygame.display.set_mode((SCREEN_WIDTH, SCREEN_HEIGHT))
sur_minigame = pygame.surface.Surface((1300, 768))
pygame.display.set_caption("Cuộc đua kì thú")
pygame.display.set_icon(pygame.image.load("SPEED RACE.png"))

# vị trí của nhân vật và các chỉ số chạy mảng
w11_x = 1
w12_x = 1
w13_x = 1
w14_x = 1
w15_x = 1
wi11 = 0
wi12 = 0
wi13 = 0
wi14 = 0
wi15 = 0

# random tốc độ nhân vật
x1 = random.uniform(0.8, 1.2)
x2 = random.uniform(0.8, 1.2)
x3 = random.uniform(0.8, 1.2)
x4 = random.uniform(0.8, 1.2)
x5 = random.uniform(0.8, 1.2)

# font chữ ở nhiều kích thước khác nhau
game_font1 = pygame.font.Font("font/DVN-Fredoka-Bold.ttf", 40)
game_font2 = pygame.font.Font("font/DVN-Fredoka-Bold.ttf", 25)
game_font3 = pygame.font.Font("font/DVN-Fredoka-Bold.ttf", 70)
game_font4 = pygame.font.Font("font/DVN-Fredoka-Bold.ttf", 80)
game_font_mini = pygame.font.Font("font/DVN-Fredoka-Bold.ttf", 10)

# luat
rule_page = 0
rule_time = 0

# buff
buff = 0

# finish game
game_finish = 0

from datetime import datetime


# Các class
def write_history(nv, result, bet, earn, nguon):
    current_date = datetime.now().date()
    formatted_date = current_date.strftime("%d/%m/%Y")
    current_time = datetime.now()
    hour = current_time.hour
    minute = current_time.minute

    time = str(hour) + "h" + str(minute) + "p" + "-" + str(formatted_date)

    stt = int(read_file("account/" + str(user) + "/stt.txt"))
    global write_history_active
    if write_history_active:
        with open('account/' + str(user) + '/stt.txt', 'w') as file:
            file.write(str(int(stt) + 1))
            write_history_active = False
    with open('account/' + str(user) + '/history.txt', 'a') as file:
        file.write(
            str(int(stt)) + ',' + str(nv) + ',' + str(result) + ',' + str(bet) + ',' + str(earn) + ',' + str(
                time) + ',' + str(nguon))
        file.write('\n')


class Background:
    def __init__(self, image):
        self.bg_img = image
        self.bg_img = pygame.transform.scale(self.bg_img, (SCREEN_WIDTH, SCREEN_HEIGHT))
        self.bg_rect = self.bg_img.get_rect(topleft=(0, 0))

    def draw_bg(self, sur):
        sur.blit(self.bg_img, self.bg_rect)


class Button:
    def __init__(self, image, x, y, w, h):
        self.but_pos = (x, y)
        self.x = x
        self.y = y
        image = pygame.transform.scale(image, (w, h))
        self.image = image
        self.image_rect = self.image.get_rect(center=(x, y))

    def draw_but(self, sur):
        sur.blit(self.image, self.image_rect)


class Text:
    def __init__(self, font, text, r, g, b, ):
        self.font = font
        self.text = text
        self.font_but = self.font.render(text, True, (r, g, b))

    def draw_text(self, sur, x, y):
        sur.blit(self.font_but, (x, y))


class nhanVat:
    def __init__(self, name, img_array, x, y, w, h):
        self.name = name
        self.img_array = img_array
        self.x = x
        self.y = y
        self.w = w
        self.h = h
        self.name = Button(img_array, x, y, w, h)
        self.quicken = 0

    def draw_nv(self, i):
        if rever[i] == 0:
            self.name.draw_but(screen)
        else:
            if rever[i] > 80:
                self.name.draw_but(screen)
                screen.blit(reverse, (self.x - 25, self.y - 25))
            else:
                screen.blit(pygame.transform.flip(self.name.image, True, False), self.name.image_rect)
            rever[i] -= 1
        if quicken_atr[i]: screen.blit(quicken, (self.x - self.w // 2, self.y))

    '''check va cham voi "sur" (box,vạch đích,...)'''

    def check_vc(self, sur):
        if self.name.image_rect.colliderect(sur.image_rect):
            return True
        else:
            return False

    def nv_img(self, i):
        self.name

    '''Cac hieu ứng khi chạm box'''

    def buff_speed(self, i):
        bonus1 = random.randrange(1, 100)
        bonus2 = random.randrange(1, 100)
        bonus3 = random.randrange(1, 100)
        bonus4 = random.randrange(1, 100)
        bonus5 = random.randrange(1, 100)

        bonus_tuple = (bonus1, bonus2, bonus3, bonus4, bonus5)
        if 1 <= bonus_tuple[i] <= 50:
            return 2
        if 50 <= bonus_tuple[i] <= 100:
            return 1.2


onevn = pygame.image.load("NPC/Vie/20.png")
twovn = pygame.image.load("NPC/Vie/21.png")
threevn = pygame.image.load("NPC/Vie/23.png")
fourvn = pygame.image.load("NPC/Vie/25.png")
fivevn = pygame.image.load("NPC/Vie/27.png")
sixvn = pygame.image.load("NPC/Vie/29.png")
A = [onevn, twovn, threevn, fourvn, fivevn, sixvn]
for i in range(0, 6, 1):
    A[i] = Background(A[i])

one = pygame.image.load("NPC/Eng/19.png")
two = pygame.image.load("NPC/Eng/22.png")
three = pygame.image.load("NPC/Eng/24.png")
four = pygame.image.load("NPC/Eng/26.png")
five = pygame.image.load("NPC/Eng/28.png")
six = pygame.image.load("NPC/Eng/30.png")
B = [one, two, three, four, five, six]
for j in range(0, 6, 1):
    B[j] = Background(B[j])

win = pygame.image.load("end/win.png")
lose = pygame.image.load("end/lose.png")
thang = pygame.image.load("end/thang.png")
thua = pygame.image.load("end/thua.png")
win_but = Button(win, 1344 // 2, 756 // 2, 800, 300)
lose_but = Button(lose, 1344 // 2, 400, 800, 400)
winvn_but = Button(thang, 1344 // 2, 400, 800, 400)
losevn_but = Button(thua, 1344 // 2, 400, 800, 400)

'''Load set 1(cho map1)'''
w11_name = ["w110", "w111", "w112", "w113", "w114", "w115", "w116", "w117", "w118", "w119", "w1110", "w1111", "w1112",
            "w1113", "w1114", "w1115", "w1116", "w1117", ]
w11_loca = ["set01/PNG1 Sequences/Walking/w110.png",
            "set01/PNG1 Sequences/Walking/w111.png",
            "set01/PNG1 Sequences/Walking/w112.png",
            "set01/PNG1 Sequences/Walking/w113.png",
            "set01/PNG1 Sequences/Walking/w114.png",
            "set01/PNG1 Sequences/Walking/w115.png",
            "set01/PNG1 Sequences/Walking/w116.png",
            "set01/PNG1 Sequences/Walking/w117.png",
            "set01/PNG1 Sequences/Walking/w118.png",
            "set01/PNG1 Sequences/Walking/w119.png",
            "set01/PNG1 Sequences/Walking/w1110.png",
            "set01/PNG1 Sequences/Walking/w1111.png",
            "set01/PNG1 Sequences/Walking/w1112.png",
            "set01/PNG1 Sequences/Walking/w1113.png",
            "set01/PNG1 Sequences/Walking/w1114.png",
            "set01/PNG1 Sequences/Walking/w1115.png",
            "set01/PNG1 Sequences/Walking/w1116.png",
            "set01/PNG1 Sequences/Walking/w1117.png"]
for i in range(0, 18):
    w11_name[i] = pygame.image.load(w11_loca[i])
for i in range(0, 18):
    w11_name[i] = pygame.transform.scale(w11_name[i], (1400, 1400))

# Load nv set1 2
w12_name = ["w120", "w121", "w122", "w123", "w124", "w125", "w126", "w127", "w128", "w129", "w1210", "w1211", "w1212",
            "w1213", "w1214", "w1215", "w1216", "w1217", ]
w12_loca = ["set01/PNG2 Sequences/Walking/w120.png",
            "set01/PNG2 Sequences/Walking/w121.png",
            "set01/PNG2 Sequences/Walking/w122.png",
            "set01/PNG2 Sequences/Walking/w123.png",
            "set01/PNG2 Sequences/Walking/w124.png",
            "set01/PNG2 Sequences/Walking/w125.png",
            "set01/PNG2 Sequences/Walking/w126.png",
            "set01/PNG2 Sequences/Walking/w127.png",
            "set01/PNG2 Sequences/Walking/w128.png",
            "set01/PNG2 Sequences/Walking/w129.png",
            "set01/PNG2 Sequences/Walking/w1210.png",
            "set01/PNG2 Sequences/Walking/w1211.png",
            "set01/PNG2 Sequences/Walking/w1212.png",
            "set01/PNG2 Sequences/Walking/w1213.png",
            "set01/PNG2 Sequences/Walking/w1214.png",
            "set01/PNG2 Sequences/Walking/w1215.png",
            "set01/PNG2 Sequences/Walking/w1216.png",
            "set01/PNG2 Sequences/Walking/w1217.png"]
for i in range(0, 18):
    w12_name[i] = pygame.image.load(w12_loca[i])
    w12_name[i] = pygame.transform.scale(w12_name[i], (140, 140))

# Load nv set1 3
w13_name = ["w130", "w131", "w132", "w133", "w134", "w135", "w136", "w137", "w138", "w139", "w1310", "w1311", "w1312",
            "w1313", "w1314", "w315", "w1316", "w1317", ]
w13_loca = ["set01/PNG3 Sequences/Walking/w130.png",
            "set01/PNG3 Sequences/Walking/w131.png",
            "set01/PNG3 Sequences/Walking/w132.png",
            'set01/PNG3 Sequences/Walking/w133.png',
            "set01/PNG3 Sequences/Walking/w134.png",
            "set01/PNG3 Sequences/Walking/w135.png",
            "set01/PNG3 Sequences/Walking/w136.png",
            "set01/PNG3 Sequences/Walking/w137.png",
            "set01/PNG3 Sequences/Walking/w138.png",
            "set01/PNG3 Sequences/Walking/w139.png",
            "set01/PNG3 Sequences/Walking/w1310.png",
            "set01/PNG3 Sequences/Walking/w1311.png",
            "set01/PNG3 Sequences/Walking/w1312.png",
            "set01/PNG3 Sequences/Walking/w1313.png",
            "set01/PNG3 Sequences/Walking/w1314.png",
            "set01/PNG3 Sequences/Walking/w1315.png",
            "set01/PNG3 Sequences/Walking/w1316.png",
            "set01/PNG3 Sequences/Walking/w1317.png"]
for i in range(0, 18):
    w13_name[i] = pygame.image.load(w13_loca[i])
    w13_name[i] = pygame.transform.scale(w13_name[i], (140, 140))

w14_name = ["w140", "w141", "w142", "w143", "w144", "w145", "w146", "w147", "w148", "w149", "w1410", "w1411", "w1412",
            "w1413", "w1414", "w1415", "w1416", "w1417", ]
w14_loca = ["set01/PNG4 Sequences/Walking/w140.png",
            "set01/PNG4 Sequences/Walking/w141.png",
            "set01/PNG4 Sequences/Walking/w142.png",
            "set01/PNG4 Sequences/Walking/w143.png",
            "set01/PNG4 Sequences/Walking/w144.png",
            "set01/PNG4 Sequences/Walking/w145.png",
            "set01/PNG4 Sequences/Walking/w146.png",
            "set01/PNG4 Sequences/Walking/w147.png",
            "set01/PNG4 Sequences/Walking/w148.png",
            "set01/PNG4 Sequences/Walking/w149.png",
            "set01/PNG4 Sequences/Walking/w1410.png",
            "set01/PNG4 Sequences/Walking/w1411.png",
            "set01/PNG4 Sequences/Walking/w1412.png",
            "set01/PNG4 Sequences/Walking/w1413.png",
            "set01/PNG4 Sequences/Walking/w1414.png",
            "set01/PNG4 Sequences/Walking/w1415.png",
            "set01/PNG4 Sequences/Walking/w1416.png",
            "set01/PNG4 Sequences/Walking/w1417.png"]
for i in range(0, 18):
    w14_name[i] = pygame.image.load(w14_loca[i])
    w14_name[i] = pygame.transform.scale(w14_name[i], (140, 140))

w15_name = ["w150", "w151", "w152", "w153", "w154", "w155", "w156", "w157", "w158", "w159", "w1510", "w1511", "w1512",
            "w1513", "w1514", "w1515", "w1516", "w1517", ]
w15_loca = ["set01/PNG5 Sequences/Walking/w150.png",
            "set01/PNG5 Sequences/Walking/w151.png",
            "set01/PNG5 Sequences/Walking/w152.png",
            "set01/PNG5 Sequences/Walking/w153.png",
            "set01/PNG5 Sequences/Walking/w154.png",
            "set01/PNG5 Sequences/Walking/w155.png",
            "set01/PNG5 Sequences/Walking/w156.png",
            "set01/PNG5 Sequences/Walking/w157.png",
            "set01/PNG5 Sequences/Walking/w158.png",
            "set01/PNG5 Sequences/Walking/w159.png",
            "set01/PNG5 Sequences/Walking/w1510.png",
            "set01/PNG5 Sequences/Walking/w1511.png",
            "set01/PNG5 Sequences/Walking/w1512.png",
            "set01/PNG5 Sequences/Walking/w1513.png",
            "set01/PNG5 Sequences/Walking/w1514.png",
            "set01/PNG5 Sequences/Walking/w1515.png",
            "set01/PNG5 Sequences/Walking/w1516.png",
            "set01/PNG5 Sequences/Walking/w1517.png"]
for i in range(0, 18):
    w15_name[i] = pygame.image.load(w15_loca[i])
    w15_name[i] = pygame.transform.scale(w15_name[i], (140, 140))

'''Load set 2 (cho map2)'''
w21_name = ["w210", "w211", "w212", "w213", "w214", "w215", "w216", "w217", "w218", "w219", "w2110", "w2111"]

w21_loca = ["set02/PNG1 Sequences/Walking/w210.png",
            "set02/PNG1 Sequences/Walking/w211.png",
            "set02/PNG1 Sequences/Walking/w212.png",
            "set02/PNG1 Sequences/Walking/w213.png",
            "set02/PNG1 Sequences/Walking/w214.png",
            "set02/PNG1 Sequences/Walking/w215.png",
            "set02/PNG1 Sequences/Walking/w216.png",
            "set02/PNG1 Sequences/Walking/w217.png",
            "set02/PNG1 Sequences/Walking/w218.png",
            "set02/PNG1 Sequences/Walking/w219.png",
            "set02/PNG1 Sequences/Walking/w2110.png",
            "set02/PNG1 Sequences/Walking/w2111.png"]

for i in range(0, 12):
    w21_name[i] = pygame.image.load(w21_loca[i])
    w21_name[i] = pygame.transform.scale(w21_name[i], (140, 140))

w22_name = ["w220", "w221", "w222", "w223", "w224", "w225", "w226", "w227", "w228", "w229", "w2210", "w2211"]

w22_loca = ["set02/PNG2 Sequences/Walking/w220.png",
            "set02/PNG2 Sequences/Walking/w221.png",
            "set02/PNG2 Sequences/Walking/w222.png",
            "set02/PNG2 Sequences/Walking/w223.png",
            "set02/PNG2 Sequences/Walking/w224.png",
            "set02/PNG2 Sequences/Walking/w225.png",
            "set02/PNG2 Sequences/Walking/w226.png",
            "set02/PNG2 Sequences/Walking/w227.png",
            "set02/PNG2 Sequences/Walking/w228.png",
            "set02/PNG2 Sequences/Walking/w229.png",
            "set02/PNG2 Sequences/Walking/w2210.png",
            "set02/PNG2 Sequences/Walking/w2211.png"]

for i in range(0, 12):
    w22_name[i] = pygame.image.load(w22_loca[i])
    w22_name[i] = pygame.transform.scale(w22_name[i], (140, 140))
w23_name = ["w230", "w231", "w232", "w233", "w234", "w235", "w236", "w237", "w238", "w239", "w2310", "w2311"]

w23_loca = ["set02/PNG3 Sequences/Walking/w230.png",
            "set02/PNG3 Sequences/Walking/w231.png",
            "set02/PNG3 Sequences/Walking/w232.png",
            "set02/PNG3 Sequences/Walking/w233.png",
            "set02/PNG3 Sequences/Walking/w234.png",
            "set02/PNG3 Sequences/Walking/w235.png",
            "set02/PNG3 Sequences/Walking/w236.png",
            "set02/PNG3 Sequences/Walking/w237.png",
            "set02/PNG3 Sequences/Walking/w238.png",
            "set02/PNG3 Sequences/Walking/w239.png",
            "set02/PNG3 Sequences/Walking/w2310.png",
            "set02/PNG3 Sequences/Walking/w2311.png"]

for i in range(0, 12):
    w23_name[i] = pygame.image.load(w23_loca[i])
    w23_name[i] = pygame.transform.scale(w23_name[i], (140, 140))

w24_name = ["w240", "w241", "w242", "w243", "w244", "w245", "w246", "w247", "w248", "w249", "w2410", "w2411"]

w24_loca = ["set02/PNG4 Sequences/Walking/w240.png",
            "set02/PNG4 Sequences/Walking/w241.png",
            "set02/PNG4 Sequences/Walking/w242.png",
            "set02/PNG4 Sequences/Walking/w243.png",
            "set02/PNG4 Sequences/Walking/w244.png",
            "set02/PNG4 Sequences/Walking/w245.png",
            "set02/PNG4 Sequences/Walking/w246.png",
            "set02/PNG4 Sequences/Walking/w247.png",
            "set02/PNG4 Sequences/Walking/w248.png",
            "set02/PNG4 Sequences/Walking/w249.png",
            "set02/PNG4 Sequences/Walking/w2410.png",
            "set02/PNG4 Sequences/Walking/w2411.png"]

for i in range(0, 12):
    w24_name[i] = pygame.image.load(w24_loca[i])
    w24_name[i] = pygame.transform.scale(w24_name[i], (140, 140))

w25_name = ["w250", "w251", "w252", "w253", "w254", "w255", "w256", "w257", "w258", "w259", "w2510", "w2511"]

w25_loca = ["set02/PNG5 Sequences/Walking/w250.png",
            "set02/PNG5 Sequences/Walking/w251.png",
            "set02/PNG5 Sequences/Walking/w252.png",
            "set02/PNG5 Sequences/Walking/w253.png",
            "set02/PNG5 Sequences/Walking/w254.png",
            "set02/PNG5 Sequences/Walking/w255.png",
            "set02/PNG5 Sequences/Walking/w256.png",
            "set02/PNG5 Sequences/Walking/w257.png",
            "set02/PNG5 Sequences/Walking/w258.png",
            "set02/PNG5 Sequences/Walking/w259.png",
            "set02/PNG5 Sequences/Walking/w2510.png",
            "set02/PNG5 Sequences/Walking/w2511.png"]

for i in range(0, 12):
    w25_name[i] = pygame.image.load(w25_loca[i])
    w25_name[i] = pygame.transform.scale(w25_name[i], (140, 140))

'''Load set 4 (cho map2)'''
w41_name = ["w410", "w411", "w412", "w413", "w414", "w415", "w416", "w417", "w418", "w419", "w4110", "w4111", "w4112",
            "w4113", "w4114", "w4115", "w4116", "w4117", ]
w41_loca = ["set04/PNG1 Sequences/Walking/w410.png",
            "set04/PNG1 Sequences/Walking/w411.png",
            "set04/PNG1 Sequences/Walking/w412.png",
            "set04/PNG1 Sequences/Walking/w413.png",
            "set04/PNG1 Sequences/Walking/w414.png",
            "set04/PNG1 Sequences/Walking/w415.png",
            "set04/PNG1 Sequences/Walking/w416.png",
            "set04/PNG1 Sequences/Walking/w417.png",
            "set04/PNG1 Sequences/Walking/w418.png",
            "set04/PNG1 Sequences/Walking/w419.png",
            "set04/PNG1 Sequences/Walking/w4110.png",
            "set04/PNG1 Sequences/Walking/w4111.png",
            "set04/PNG1 Sequences/Walking/w4112.png",
            "set04/PNG1 Sequences/Walking/w4113.png",
            "set04/PNG1 Sequences/Walking/w4114.png",
            "set04/PNG1 Sequences/Walking/w4115.png",
            "set04/PNG1 Sequences/Walking/w4116.png",
            "set04/PNG1 Sequences/Walking/w4117.png"]
for i in range(0, 18):
    w41_name[i] = pygame.image.load(w41_loca[i])
    w41_name[i] = pygame.transform.scale(w41_name[i], (140, 140))

# Load nv set1 2
w42_name = ["w420", "w421", "w422", "w423", "w424", "w425", "w426", "w427", "w428", "w429", "w4210", "w4211", "w4212",
            "w4213", "w4214", "w4215", "w4216", "w4217", ]
w42_loca = ["set04/PNG2 Sequences/Walking/w420.png",
            "set04/PNG2 Sequences/Walking/w421.png",
            "set04/PNG2 Sequences/Walking/w422.png",
            "set04/PNG2 Sequences/Walking/w423.png",
            "set04/PNG2 Sequences/Walking/w424.png",
            "set04/PNG2 Sequences/Walking/w425.png",
            "set04/PNG2 Sequences/Walking/w426.png",
            "set04/PNG2 Sequences/Walking/w427.png",
            "set04/PNG2 Sequences/Walking/w428.png",
            "set04/PNG2 Sequences/Walking/w429.png",
            "set04/PNG2 Sequences/Walking/w4210.png",
            "set04/PNG2 Sequences/Walking/w4211.png",
            "set04/PNG2 Sequences/Walking/w4212.png",
            "set04/PNG2 Sequences/Walking/w4213.png",
            "set04/PNG2 Sequences/Walking/w4214.png",
            "set04/PNG2 Sequences/Walking/w4215.png",
            "set04/PNG2 Sequences/Walking/w4216.png",
            "set04/PNG2 Sequences/Walking/w4217.png"]
for i in range(0, 18):
    w42_name[i] = pygame.image.load(w42_loca[i])
    w42_name[i] = pygame.transform.scale(w42_name[i], (140, 140))

# Load nv set1 3
w43_name = ["w430", "w431", "w432", "w433", "w434", "w435", "w436", "w437", "w438", "w439", "w4310", "w4311", "w4312",
            "w4313", "w4314", "w315", "w4316", "w4317", ]
w43_loca = ["set04/PNG3 Sequences/Walking/w430.png",
            "set04/PNG3 Sequences/Walking/w431.png",
            "set04/PNG3 Sequences/Walking/w432.png",
            "set04/PNG3 Sequences/Walking/w433.png",
            "set04/PNG3 Sequences/Walking/w434.png",
            "set04/PNG3 Sequences/Walking/w435.png",
            "set04/PNG3 Sequences/Walking/w436.png",
            "set04/PNG3 Sequences/Walking/w437.png",
            "set04/PNG3 Sequences/Walking/w438.png",
            "set04/PNG3 Sequences/Walking/w439.png",
            "set04/PNG3 Sequences/Walking/w4310.png",
            "set04/PNG3 Sequences/Walking/w4311.png",
            "set04/PNG3 Sequences/Walking/w4312.png",
            "set04/PNG3 Sequences/Walking/w4313.png",
            "set04/PNG3 Sequences/Walking/w4314.png",
            "set04/PNG3 Sequences/Walking/w4315.png",
            "set04/PNG3 Sequences/Walking/w4316.png",
            "set04/PNG3 Sequences/Walking/w4317.png"]
for i in range(0, 18):
    w43_name[i] = pygame.image.load(w43_loca[i])
    w43_name[i] = pygame.transform.scale(w43_name[i], (140, 140))

w44_name = ["w440", "w441", "w442", "w443", "w444", "w445", "w446", "w447", "w448", "w449", "w4410", "w4411", "w4412",
            "w4413", "w4414", "w4415", "w4416", "w4417", ]
w44_loca = ["set04/PNG4 Sequences/Walking/w440.png",
            "set04/PNG4 Sequences/Walking/w441.png",
            "set04/PNG4 Sequences/Walking/w442.png",
            "set04/PNG4 Sequences/Walking/w443.png",
            "set04/PNG4 Sequences/Walking/w444.png",
            "set04/PNG4 Sequences/Walking/w445.png",
            "set04/PNG4 Sequences/Walking/w446.png",
            "set04/PNG4 Sequences/Walking/w447.png",
            "set04/PNG4 Sequences/Walking/w448.png",
            "set04/PNG4 Sequences/Walking/w449.png",
            "set04/PNG4 Sequences/Walking/w4410.png",
            "set04/PNG4 Sequences/Walking/w4411.png",
            "set04/PNG4 Sequences/Walking/w4412.png",
            "set04/PNG4 Sequences/Walking/w4413.png",
            "set04/PNG4 Sequences/Walking/w4414.png",
            "set04/PNG4 Sequences/Walking/w4415.png",
            "set04/PNG4 Sequences/Walking/w4416.png",
            "set04/PNG4 Sequences/Walking/w4417.png"]
for i in range(0, 18):
    w44_name[i] = pygame.image.load(w44_loca[i])
    w44_name[i] = pygame.transform.scale(w44_name[i], (140, 140))

w45_name = ["w450", "w451", "w452", "w453", "w454", "w455", "w456", "w457", "w458", "w459", "w4510", "w4511", "w4512",
            "w4513", "w4514", "w4515", "w4516", "w4517", ]
w45_loca = ["set04/PNG5 Sequences/Walking/w450.png",
            "set04/PNG5 Sequences/Walking/w451.png",
            "set04/PNG5 Sequences/Walking/w452.png",
            "set04/PNG5 Sequences/Walking/w453.png",
            "set04/PNG5 Sequences/Walking/w454.png",
            "set04/PNG5 Sequences/Walking/w455.png",
            "set04/PNG5 Sequences/Walking/w456.png",
            "set04/PNG5 Sequences/Walking/w457.png",
            "set04/PNG5 Sequences/Walking/w458.png",
            "set04/PNG5 Sequences/Walking/w459.png",
            "set04/PNG5 Sequences/Walking/w4510.png",
            "set04/PNG5 Sequences/Walking/w4511.png",
            "set04/PNG5 Sequences/Walking/w4512.png",
            "set04/PNG5 Sequences/Walking/w4513.png",
            "set04/PNG5 Sequences/Walking/w4514.png",
            "set04/PNG5 Sequences/Walking/w4515.png",
            "set04/PNG5 Sequences/Walking/w4516.png",
            "set04/PNG5 Sequences/Walking/w4517.png"]
for i in range(0, 18):
    w45_name[i] = pygame.image.load(w45_loca[i])
    w45_name[i] = pygame.transform.scale(w45_name[i], (140, 140))

'''Load set 5 (cho map2)'''
w51_name = ["w5101", "w5102", "w5103", "w5104", "w5105", "w5106", "w5107", "w5108", "w5109", "w5110", "w5111", "w5112",
            "w5113", "w5114", "w5115", "w5116", "w5117", "w518"]
w51_loca = [
    "set05/PNG1 Sequences/Walking/w5101.png",
    "set05/PNG1 Sequences/Walking/w5102.png",
    "set05/PNG1 Sequences/Walking/w5103.png",
    "set05/PNG1 Sequences/Walking/w5104.png",
    "set05/PNG1 Sequences/Walking/w5105.png",
    "set05/PNG1 Sequences/Walking/w5106.png",
    "set05/PNG1 Sequences/Walking/w5107.png",
    "set05/PNG1 Sequences/Walking/w5108.png",
    "set05/PNG1 Sequences/Walking/w5109.png",
    "set05/PNG1 Sequences/Walking/w5110.png",
    "set05/PNG1 Sequences/Walking/w5111.png",
    "set05/PNG1 Sequences/Walking/w5112.png",
    "set05/PNG1 Sequences/Walking/w5113.png",
    "set05/PNG1 Sequences/Walking/w5114.png",
    "set05/PNG1 Sequences/Walking/w5115.png",
    "set05/PNG1 Sequences/Walking/w5116.png",
    "set05/PNG1 Sequences/Walking/w5117.png",
    "set05/PNG1 Sequences/Walking/w5118.png", ]
for i in range(0, 18):
    w51_name[i] = pygame.image.load(w51_loca[i])
    w51_name[i] = pygame.transform.scale(w51_name[i], (120, 120))
w52_name = ["w5201", "w5202", "w5203", "w5204", "w5205", "w5206", "w5207", "w5208", "w5209", "w5210", "w5211", "w5212",
            "w5213", "w5214", "w5215", "w5216", "w5217", "w5218"]
w52_loca = [
    "set05/PNG2 Sequences/Walking/w5201.png",
    "set05/PNG2 Sequences/Walking/w5202.png",
    "set05/PNG2 Sequences/Walking/w5203.png",
    "set05/PNG2 Sequences/Walking/w5204.png",
    "set05/PNG2 Sequences/Walking/w5205.png",
    "set05/PNG2 Sequences/Walking/w5206.png",
    "set05/PNG2 Sequences/Walking/w5207.png",
    "set05/PNG2 Sequences/Walking/w5208.png",
    "set05/PNG2 Sequences/Walking/w5209.png",
    "set05/PNG2 Sequences/Walking/w5210.png",
    "set05/PNG2 Sequences/Walking/w5211.png",
    "set05/PNG2 Sequences/Walking/w5212.png",
    "set05/PNG2 Sequences/Walking/w5213.png",
    "set05/PNG2 Sequences/Walking/w5214.png",
    "set05/PNG2 Sequences/Walking/w5215.png",
    "set05/PNG2 Sequences/Walking/w5216.png",
    "set05/PNG2 Sequences/Walking/w5217.png",
    "set05/PNG2 Sequences/Walking/w5218.png"]
for i in range(0, 18):
    w52_name[i] = pygame.image.load(w52_loca[i])
    w52_name[i] = pygame.transform.scale(w52_name[i], (120, 120))

w53_name = ["w5301", "w5302", "w5303", "w5304", "w5305", "w5306", "w5307", "w5308", "w5309", "w5310", "w5311", "w5312",
            "w5313", "w5314", "w5315", "w5316", "w5317", "w5318"]
w53_loca = [
    "set05/PNG3 Sequences/Walking/w5301.png",
    "set05/PNG3 Sequences/Walking/w5302.png",
    "set05/PNG3 Sequences/Walking/w5303.png",
    "set05/PNG3 Sequences/Walking/w5304.png",
    "set05/PNG3 Sequences/Walking/w5305.png",
    "set05/PNG3 Sequences/Walking/w5306.png",
    "set05/PNG3 Sequences/Walking/w5307.png",
    "set05/PNG3 Sequences/Walking/w5308.png",
    "set05/PNG3 Sequences/Walking/w5309.png",
    "set05/PNG3 Sequences/Walking/w5310.png",
    "set05/PNG3 Sequences/Walking/w5311.png",
    "set05/PNG3 Sequences/Walking/w5312.png",
    "set05/PNG3 Sequences/Walking/w5313.png",
    "set05/PNG3 Sequences/Walking/w5314.png",
    "set05/PNG3 Sequences/Walking/w5315.png",
    "set05/PNG3 Sequences/Walking/w5316.png",
    "set05/PNG3 Sequences/Walking/w5317.png",
    "set05/PNG3 Sequences/Walking/w5318.png"]
for i in range(0, 18):
    w53_name[i] = pygame.image.load(w53_loca[i])
    w53_name[i] = pygame.transform.scale(w53_name[i], (120, 120))
w54_name = ["w5401", "w5402", "w5403", "w5404", "w5405", "w5406", "w5407", "w5408", "w5409", "w5410", "w5411", "w5412",
            "w5413", "w5414", "w5415", "w5416", "w5417", "w5418"]
w54_loca = [
    "set05/PNG4 Sequences/Walking/w5401.png",
    "set05/PNG4 Sequences/Walking/w5402.png",
    "set05/PNG4 Sequences/Walking/w5403.png",
    "set05/PNG4 Sequences/Walking/w5404.png",
    "set05/PNG4 Sequences/Walking/w5405.png",
    "set05/PNG4 Sequences/Walking/w5406.png",
    "set05/PNG4 Sequences/Walking/w5407.png",
    "set05/PNG4 Sequences/Walking/w5408.png",
    "set05/PNG4 Sequences/Walking/w5409.png",
    "set05/PNG4 Sequences/Walking/w5410.png",
    "set05/PNG4 Sequences/Walking/w5411.png",
    "set05/PNG4 Sequences/Walking/w5412.png",
    "set05/PNG4 Sequences/Walking/w5413.png",
    "set05/PNG4 Sequences/Walking/w5414.png",
    "set05/PNG4 Sequences/Walking/w5415.png",
    "set05/PNG4 Sequences/Walking/w5416.png",
    "set05/PNG4 Sequences/Walking/w5417.png",
    "set05/PNG4 Sequences/Walking/w5418.png"]
for i in range(0, 18):
    w54_name[i] = pygame.image.load(w54_loca[i])
    w54_name[i] = pygame.transform.scale(w54_name[i], (120, 120))
w55_name = ["w5501", "w5502", "w5503", "w5504", "w5505", "w5506", "w5507", "w5508", "w5509", "w5510", "w5511", "w5512",
            "w5513", "w5514", "w5515", "w5516", "w5517", "w5518"]
w55_loca = [
    "set05/PNG5 Sequences/Walking/w5501.png",
    "set05/PNG5 Sequences/Walking/w5502.png",
    "set05/PNG5 Sequences/Walking/w5503.png",
    "set05/PNG5 Sequences/Walking/w5504.png",
    "set05/PNG5 Sequences/Walking/w5505.png",
    "set05/PNG5 Sequences/Walking/w5506.png",
    "set05/PNG5 Sequences/Walking/w5507.png",
    "set05/PNG5 Sequences/Walking/w5508.png",
    "set05/PNG5 Sequences/Walking/w5509.png",
    "set05/PNG5 Sequences/Walking/w5510.png",
    "set05/PNG5 Sequences/Walking/w5511.png",
    "set05/PNG5 Sequences/Walking/w5512.png",
    "set05/PNG5 Sequences/Walking/w5513.png",
    "set05/PNG5 Sequences/Walking/w5514.png",
    "set05/PNG5 Sequences/Walking/w5515.png",
    "set05/PNG5 Sequences/Walking/w5516.png",
    "set05/PNG5 Sequences/Walking/w5517.png",
    "set05/PNG5 Sequences/Walking/w5518.png"]
for i in range(0, 18):
    w55_name[i] = pygame.image.load(w55_loca[i])
    w55_name[i] = pygame.transform.scale(w55_name[i], (120, 120))

# Background

'''menu chính'''
bgvn = pygame.image.load("background/menugame.png")
bg = pygame.image.load("background/menugamevn.png")
background = Background(bg)
backgroundvn = Background(bgvn)

'''màn hình chọn nhân vật'''
bg0 = pygame.image.load("background/menugamemap.png")
background0 = Background(bg0)

'''các level của map1'''
bg1 = pygame.image.load("background/1.jpg")
bg1_lv2 = pygame.image.load("background/1.2.jpg")
bg1_lv3 = pygame.image.load("background/1.3.jpg")
background1 = Background(bg1)
background1_lv2 = Background(bg1_lv2)
background1_lv3 = Background(bg1_lv3)
'''các level của map2'''
bg2 = pygame.image.load("background/2.jpg")
bg2_lv2 = pygame.image.load("background/2.2.jpg")
bg2_lv3 = pygame.image.load("background/2.3.jpg")
background2 = Background(bg2)
background2_lv2 = Background(bg2_lv2)
background2_lv3 = Background(bg2_lv3)

'''các level của map3'''
bg3 = pygame.image.load("background/3.jpg")
bg3_lv2 = pygame.image.load("background/3.2.jpg")
bg3_lv3 = pygame.image.load("background/3.3.jpg")
background3 = Background(bg3)
background3_lv2 = Background(bg3_lv2)
background3_lv3 = Background(bg3_lv3)
'''kết quả'''
result = pygame.image.load("background/result.png")
background6 = Background(result)

'''luật chơi'''
rule_bg = pygame.image.load("background/rulebg.png")
background7 = Background(rule_bg)

# button
'''nút play'''
play = pygame.image.load("Button/play.png")
play_but = Button(play, 450, 580, 105, 105)
next_img = pygame.image.load("Button/next.png")

fp = pygame.image.load("Button/fp.png")
sn = pygame.image.load("Button/sn.png")
fr = pygame.image.load("Button/fruit.png")

fp_but = Button(fp, 200, 400, 300, 300)
fr_but = Button(fr, 660, 400, 300, 300)
sn_but = Button(sn, 1100, 400, 300, 300)

choose1 = pygame.image.load("Button/set_but1.png")
choose1vn = pygame.image.load("Button/set_but1vn.png")

choose2 = pygame.image.load("Button/set_but2.png")
choose2vn = pygame.image.load("Button/set_but2vn.png")

choose3 = pygame.image.load("Button/set_but3.png")
choose3vn = pygame.image.load("Button/set_but3vn.png")
set1vn_but = Button(choose1vn, set_but_x, set_but_y, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 4)
set2vn_but = Button(choose2vn, set_but_x, set_but_y + set_but_dis, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 4)
set3vn_but = Button(choose3vn, set_but_x, set_but_y + 2 * set_but_dis, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 4)

set1_but = Button(choose1, set_but_x, set_but_y, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 4)
set2_but = Button(choose2, set_but_x, set_but_y + set_but_dis, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 4)
set3_but = Button(choose3, set_but_x, set_but_y + 2 * set_but_dis, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 4)

exit = pygame.image.load("Button/exit.png")
exit_but = Button(exit, exit_but_x, exit_but_y, exit_w, exit_h)

signup = pygame.image.load("Button/signup.png")
signupvn = pygame.image.load("Button/signupvn.png")
login = pygame.image.load("Button/login.png")
loginvn = pygame.image.load("Button/loginvn.png")
logout = pygame.image.load("Button/logout.png")
logoutvn = pygame.image.load("Button/logoutvn.png")
signup_but = Button(signup, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 1.3, 250, 130)
login_but = Button(login, SCREEN_WIDTH // 2.8, SCREEN_HEIGHT // 1.3, 250, 130)
signupvn_but = Button(signupvn, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 1.3, 250, 130)
logout_but = Button(logout, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 4)
loginvn_but = Button(loginvn, SCREEN_WIDTH // 2.8, SCREEN_HEIGHT // 1.3, 250, 130)
logoutvn_but = Button(logoutvn, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 4)

setting = pygame.image.load("Button/setting.png")
setting_but = Button(setting, SCREEN_WIDTH // 1.025, SCREEN_HEIGHT // 17, SCREEN_WIDTH // 20, SCREEN_HEIGHT // 12)

language = pygame.image.load("Button/language.png")
languagevn = pygame.image.load("Button/languagevn.png")
vie = pygame.image.load("Button/vie.png")
eng = pygame.image.load("Button/eng.png")
language_but = Button(language, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 1.5, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 4)
languagevn_but = Button(languagevn, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 1.5, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 4)
vie_but = Button(vie, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 2, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 4)
eng_but = Button(eng, SCREEN_WIDTH // 5, SCREEN_HEIGHT // 1.5, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 4)

box = pygame.image.load("Button/box.png")

car1 = pygame.image.load("setbg3/set03/car1.png")
car2 = pygame.image.load("setbg3/set03/car2.png")
car3 = pygame.image.load("setbg3/set03/car3.png")
car4 = pygame.image.load("setbg3/set03/car4a.png")
car5 = pygame.image.load("setbg3/set03/car5.png")

banhxe = pygame.image.load("setbg3/set03/banhxe2.png")

choose_set1 = pygame.image.load("Button/set1.png")
choose_set1vn = pygame.image.load("Button/set1vn.png")
choose_set2 = pygame.image.load("Button/set2.png")
choose_set2vn = pygame.image.load("Button/set2vn.png")
choose_set1_but = Button(choose_set1, SCREEN_WIDTH // 2, 50, 250, 130)
choose_set1vn_but = Button(choose_set1vn, SCREEN_WIDTH // 2, 50, 250, 130)
choose_set2_but = Button(choose_set2, SCREEN_WIDTH // 5, 50, 250, 130)
choose_set2vn_but = Button(choose_set1vn, SCREEN_WIDTH // 5, 50, 250, 130)
minigame = pygame.image.load("Button/minigame.png")
minigame_but = Button(minigame, 300, 712, 160, 60)
minigamevn = pygame.image.load("Button/minigamevn.png")
minigamevn_but = Button(minigamevn, 300, 712, 160, 60)

snake = pygame.image.load("Button/snake.png")
fruit = pygame.image.load("Button/fruit.png")
flappy_bird = pygame.image.load("Button/bird.png")
flappy_bird_but = Button(flappy_bird, 200, 600, SCREEN_WIDTH // 4, SCREEN_HEIGHT // 15)
fruit_but = Button(fruit, 660, 600, SCREEN_WIDTH // 10, SCREEN_HEIGHT // 15)
snake_but = Button(snake, 1100, 600, SCREEN_WIDTH // 10, SCREEN_HEIGHT // 15)

rule = pygame.image.load("Button/rule.png")
rulevn = pygame.image.load("Button/rulevn.png")
rule_info = pygame.image.load("Button/rules.png")
rule_info = pygame.transform.scale(rule_info, (1920 / 2, 1080 / 2))
rule_but = Button(rule, 480, 712, 100, 55)
rulevn_but = Button(rulevn, 480, 718, 120, 60)

coin_img = pygame.image.load("Button/coin.png")
coin_but = Button(coin_img, 80, 60, 50, 50)

demo_map1_but = Button(bg1, demomap_x, demomap_y, demomap_w, demomap_h)
demo_map2_but = Button(bg2, demomap_x, demomap_y, demomap_w, demomap_h)
demo_map3_but = Button(bg3, demomap_x, demomap_y, demomap_w, demomap_h)
board = pygame.image.load("Button/board.png")
board_but = Button(board, 1240, 90, 200, 180)

'''Chọn nhân vật'''

Nv11 = pygame.image.load("set01/PNG1.png")
Nv12 = pygame.image.load("set01/PNG2.png")
Nv13 = pygame.image.load("set01/PNG3.png")
Nv14 = pygame.image.load("set01/PNG4.png")
Nv15 = pygame.image.load("set01/PNG5.png")

Nv11_but = Button(Nv11, 400, 200, 150, 175)
Nv12_but = Button(Nv12, 600, 200, 150, 175)
Nv13_but = Button(Nv13, 800, 200, 150, 175)
Nv14_but = Button(Nv14, 1000, 200, 150, 175)
Nv15_but = Button(Nv15, 1200, 200, 150, 175)

Nv11_but_clicked = Button(Nv11, 750, 400, 150 * 1.5, 175 * 1.5)
Nv12_but_clicked = Button(Nv12, 750, 400, 150 * 1.5, 175 * 1.5)
Nv13_but_clicked = Button(Nv13, 750, 400, 150 * 1.5, 175 * 1.5)
Nv14_but_clicked = Button(Nv14, 750, 400, 150 * 1.5, 175 * 1.5)
Nv15_but_clicked = Button(Nv15, 750, 400, 150 * 1.5, 175 * 1.5)

Nv41 = pygame.image.load("set04/nv1.png")
Nv42 = pygame.image.load("set04/nv2.png")
Nv43 = pygame.image.load("set04/nv3.png")
Nv44 = pygame.image.load("set04/nv4.png")
Nv45 = pygame.image.load("set04/nv5.png")

Nv41_but = Button(Nv41, 400, 200, 150, 175)
Nv42_but = Button(Nv42, 600, 200, 150, 175)
Nv43_but = Button(Nv43, 800, 200, 150, 175)
Nv44_but = Button(Nv44, 1000, 200, 150, 175)
Nv45_but = Button(Nv45, 1200, 200, 150, 175)

Nv21 = pygame.image.load("set02/nv1.png")
Nv22 = pygame.image.load("set02/nv2.png")
Nv23 = pygame.image.load("set02/nv3.png")
Nv24 = pygame.image.load("set02/nv4.png")
Nv25 = pygame.image.load("set02/nv5.png")

Nv21_but = Button(Nv21, 400, 200, 150, 175)
Nv22_but = Button(Nv22, 600, 200, 150, 175)
Nv23_but = Button(Nv23, 800, 200, 150, 175)
Nv24_but = Button(Nv24, 1000, 200, 150, 175)
Nv25_but = Button(Nv25, 1200, 200, 150, 175)

Nv51 = pygame.image.load("set05/nv1.png")
Nv52 = pygame.image.load("set05/nv2.png")
Nv53 = pygame.image.load("set05/nv3.png")
Nv54 = pygame.image.load("set05/nv4.png")
Nv55 = pygame.image.load("set05/nv5.png")

Nv51_but = Button(Nv51, 400, 200, 150, 175)
Nv52_but = Button(Nv52, 600, 200, 150, 175)
Nv53_but = Button(Nv53, 800, 200, 150, 175)
Nv54_but = Button(Nv54, 1000, 200, 150, 175)
Nv55_but = Button(Nv55, 1200, 200, 150, 175)

Nv31 = pygame.image.load("set03/nhu.png")
Nv32 = pygame.image.load("set03/phong.png")
Nv33 = pygame.image.load("set03/phu.png")
Nv34 = pygame.image.load("set03/thong.png")
Nv35 = pygame.image.load("set03/quoc_thanh.png")

Nv31_but = Button(Nv31, 400, 200, 150, 175)
Nv32_but = Button(Nv32, 600, 200, 150, 175)
Nv33_but = Button(Nv33, 800, 200, 150, 175)
Nv34_but = Button(Nv34, 1000, 200, 150, 175)
Nv35_but = Button(Nv35, 1200, 200, 150, 175)

# Head
head11 = pygame.image.load("head/Dau nv/Set1/PNG1.png")
head12 = pygame.image.load("head/Dau nv/Set1/PNG2.png")
head13 = pygame.image.load("head/Dau nv/Set1/PNG3.png")
head14 = pygame.image.load("head/Dau nv/Set1/PNG4.png")
head15 = pygame.image.load("head/Dau nv/Set1/PNG5.png")

'''
head11=pygame.transform.scale(head11,(210,210))
head12=pygame.transform.scale(head12,(280,280))
head13=pygame.transform.scale(head13,(200,200))
head14=pygame.transform.scale(head14,(200,200))
head15=pygame.transform.scale(head15,(190,200))'''

head21 = pygame.image.load("head/Dau nv/Set 2/3.png")
head22 = pygame.image.load("head/Dau nv/Set 2/4.png")
head23 = pygame.image.load("head/Dau nv/Set 2/2.png")
head24 = pygame.image.load("head/Dau nv/Set 2/1.png")
head25 = pygame.image.load("head/Dau nv/Set 2/5.png")

head31 = pygame.image.load("head/Dau nv/Set3/nv1.png")
head32 = pygame.image.load("head/Dau nv/Set3/nv2.png")
head33 = pygame.image.load("head/Dau nv/Set3/nv3.png")
head34 = pygame.image.load("head/Dau nv/Set3/nv4.png")
head35 = pygame.image.load("head/Dau nv/Set3/nv5.png")

head41 = pygame.image.load("head/Dau nv/Set4/nv1.png")
head42 = pygame.image.load("head/Dau nv/Set4/nv2.png")
head43 = pygame.image.load("head/Dau nv/Set4/nv3.png")
head44 = pygame.image.load("head/Dau nv/Set4/nv4.png")
head45 = pygame.image.load("head/Dau nv/Set4/nv5.png")

head51 = pygame.image.load("head/Dau nv/Set5/nv1.png")
head52 = pygame.image.load("head/Dau nv/Set5/nv2.png")
head53 = pygame.image.load("head/Dau nv/Set5/nv3.png")
head54 = pygame.image.load("head/Dau nv/Set5/nv4.png")
head55 = pygame.image.load("head/Dau nv/Set5/nv5.png")

history = pygame.image.load("Button/history.png")
history_but = Button(history, 650, 720, 150, 65)
historyvn = pygame.image.load("Button/historyvn.png")
historyvn_but = Button(historyvn, 650, 718, 150, 50)

# Set text
chooseMap_text = Text(game_font1, "Choose the map", 255, 255, 255)
chooseMapvn_text = Text(game_font1, "Chọn bản đồ", 255, 255, 255)
set1_text = Text(game_font4, "SET 1", 0, 0, 0)
set1vn_text = Text(game_font4, "Nhóm 1", 0, 0, 0)

set2_text = Text(game_font4, "SET 2", 0, 0, 0)
set2vn_text = Text(game_font4, "Nhóm 2", 0, 0, 0)
choosenv = Text(game_font3, "Choose character", 0, 0, 0)
choosenv_vn = Text(game_font3, "Chọn nhân vật", 0, 0, 0)

Size1 = Text(game_font1, "100%", 0, 0, 0)
size1_active = Text(game_font1, "100%", 255, 0, 0)
Size2 = Text(game_font1, "80%", 0, 0, 0)
size2_active = Text(game_font1, "80%", 255, 0, 0)
Size3 = Text(game_font1, "60%", 0, 0, 0)
size3_active = Text(game_font1, "60%", 255, 0, 0)

# Quang duong nhan vat di duoc
S_nv1 = Text(game_font2, str(int(w11_x / 3)), 0, 0, 0)
S_nv2 = Text(game_font2, str(int(w12_x / 3)), 0, 0, 0, )
S_nv3 = Text(game_font2, str(int(w13_x / 3)), 0, 0, 0)
S_nv4 = Text(game_font2, str(int(w14_x / 3)), 0, 0, 0)
S_nv5 = Text(game_font2, str(int(w15_x / 3)), 0, 0, 0)


# Các hàm
def check_press(rect, pos):
    if rect.collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            pygame.mixer.init()
            click_music = pygame.mixer.Sound("music/click.mp3")
            click_music.play()
            time.sleep(0.1)
            return True
        else:
            return False


def check_hover(rect, pos):
    if rect.collidepoint(pos):
        return True
    else:
        return False


'''Tắt hết các màn hình còn lại'''


def off_screen_except(x):
    global set_, wait, set1, set2, set3, setting_bool, result, lang, start_bool, minigame_bool, rule_bool, choosenv_bool1
    global setnv11, setnv21, choosenv_bool2, setnv12, setnv22, flap, chonnv1, chonnv2, choosenv_bool3, ketqua, history_bool
    set_ = False  # menu
    wait = False  # chon map
    set1 = False  # map1
    set2 = False  # map2
    set3 = False  # map3
    result = False  # ket qua
    lang = False  # ngon ngu
    start_bool = False  # bat dau tran
    setting_bool = False  # setting
    minigame_bool = False  # minigame
    rule_bool = False
    flap = False  # game flappy bird
    choosenv_bool1 = choosenv_bool2 = False  # chon set truoc khi vao
    setnv11 = False  # set 1 của map1
    setnv21 = False  # set 2 của map1
    setnv12 = False  # set 1 của map2
    setnv22 = False  # set 2 của map2
    chonnv1 = False
    chonnv2 = False
    choosenv_bool3 = False
    ketqua = False
    history_bool = False

    if x == 0: set_ = True
    if x == 1: set1 = True
    if x == 2: set2 = True
    if x == 3: set3 = True
    if x == 4: set4 = True  # đã bỏ không xài
    if x == 5: set5 = True  # đã bỏ không xài
    if x == 6: setting_bool = True
    if x == 7: lang = True
    if x == 8: result = True
    if x == 9: start_bool = True
    if x == 10: minigame_bool = True
    if x == 11: rule_bool = True
    if x == 12: choosenv_bool1 = True  # chon nv map 1
    if x == 15: choosenv_bool2 = True  # chon nv map2
    if x == 16: flap = True
    if x == 21: choosenv_bool3 = True
    if x == 22: ketqua = True
    if x == 23: history_bool = True
    if x == 0.5: wait = True


def language_screen():
    background0.draw_bg(screen)

    off_screen_except(7)

    vie_but.draw_but(screen)
    eng_but.draw_but(screen)

    if check_press(vie_but.image_rect, pos):
        with open('lan.txt', 'w') as file:
            file.write(str(0))
        off_screen_except(0)

    if check_press(eng_but.image_rect, pos):
        with open('lan.txt', 'w') as file:
            file.write(str(1))
        off_screen_except(0)

    if check_press(exit_but.image_rect, pos):
        off_screen_except(6)
        time.sleep(0.1)


face_id_img = pygame.image.load("Button/face_id.png")
face_id_but = Button(face_id_img, SCREEN_WIDTH // 1.95, SCREEN_HEIGHT // 1.3, 250, 130)


def account_login():
    if lock:
        background.draw_bg(screen)
        play_but.draw_but(screen)
        setting_but.draw_but(screen)

    else:
        background.draw_bg(screen)
        if lan:
            signup_but.draw_but(screen)
            login_but.draw_but(screen)

        else:
            signupvn_but.draw_but(screen)
            loginvn_but.draw_but(screen)

        # Chay file dang nhap
        if check_press(signup_but.image_rect, pos):
            subprocess.run(["python", "RegistrationApp.py"])
        if check_press(login_but.image_rect, pos):
            subprocess.run(["python", "LoginApp.py"])


r = g = b = 255


def menu_screen():
    global background1, background2, background3
    background1 = Background(bg1)
    background2 = Background(bg2)
    background3 = Background(bg3)
    win_but.draw_but(screen)
    global r, g, b
    global get_coin  # check xem co cong coin khi thang chua
    get_coin = False

    if set_ and lock == False:
        account_login()

    if lock:

        if lan:
            background.draw_bg(screen)

            minigame_but.draw_but(screen)
            history_but.draw_but(screen)
            rule_but.draw_but(screen)

        else:
            backgroundvn.draw_bg(screen)
            minigamevn_but.draw_but(screen)
            historyvn_but.draw_but(screen)
            rulevn_but.draw_but(screen)
        play_but.draw_but(screen)

        setting_but.draw_but(screen)

        if check_press(history_but.image_rect, pos):
            off_screen_except(23)

        if check_press(play_but.image_rect, pos):
            off_screen_except(0.5)

        if check_press(setting_but.image_rect, pos):
            off_screen_except(6)

        if check_press(minigame_but.image_rect, pos):
            off_screen_except(10)
            minigame_screen()
        if check_press(rule_but.image_rect, pos):
            off_screen_except(11)

        coin_text = Text(game_font1, str(coin), 255, 200, 60)
        coin_text.draw_text(screen, 10, 35)
        coin_but.draw_but(screen)
        user_text = Text(game_font1, str(user), 255, 200, 60)
        user_text.draw_text(screen, 10, 2)
    else:
        background.draw_bg(screen)
        face_id_but.draw_but(screen)
        Vie_text = Text(game_font1, "VIE", abs(255 - r), abs(255 - g), abs(255 - b))
        Eng_text = Text(game_font1, "ENG", r, g, b)
        Vie_text.draw_text(screen, 100, 710)
        Eng_text.draw_text(screen, 180, 710)
        if pygame.rect.Rect(100, 710, 60, 30).collidepoint(pos):
            if pygame.mouse.get_pressed()[0] == 1:
                r = g = b = 0
                with open('lan.txt', 'w') as file:
                    file.write("0")
        if pygame.rect.Rect(180, 710, 60, 30).collidepoint(pos):
            if pygame.mouse.get_pressed()[0] == 1:
                r = g = b = 255
                with open('lan.txt', 'w') as file:
                    file.write("1")
        if check_press(face_id_but.image_rect, pos):
            log, account = function_faceid.FaceID(1, 1344, 756, screen)
            if log:
                with open('log.txt', 'w') as file:
                    file.write(str(1))
            with open('user.txt', 'w') as file:
                file.write(str(account))

        if lan:
            signup_but.draw_but(screen)
            login_but.draw_but(screen)

        else:
            signupvn_but.draw_but(screen)
            loginvn_but.draw_but(screen)
        setting_but.draw_but(screen)

    global w11_x, w12_x, w13_x, w14_x, w15_x, wi11, wi12, wi13, wi14, wi15
    w11_x = 1
    w12_x = 1
    w13_x = 1
    w14_x = 1
    w15_x = 1
    wi11 = 0
    wi12 = 0
    wi13 = 0
    wi14 = 0
    wi15 = 0


def minigame_screen():
    global screen, coin
    off_screen_except(10)
    background0.draw_bg(screen)
    fr_but.draw_but(screen)
    sn_but.draw_but(screen)
    fp_but.draw_but(screen)

    if check_press(fp_but.image_rect, pos):
        coin = int(PlappyBird.play(coin))
        screen = pygame.display.set_mode((SCREEN_WIDTH, SCREEN_HEIGHT))

    if check_press(sn_but.image_rect, pos):
        coin = int(Snake_game_class.play(coin))
        screen = pygame.display.set_mode((SCREEN_WIDTH, SCREEN_HEIGHT))

    if check_press(fr_but.image_rect, pos):
        coin = int(Fruit_ninja.play(coin))
        screen = pygame.display.set_mode((SCREEN_WIDTH, SCREEN_HEIGHT))

    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0)


def wait_screen():
    global bet
    bet = ""
    background0.draw_bg(screen)

    if lan:
        chooseMap_text.draw_text(screen, choose_map_x, choose_map_y)
    else:
        chooseMapvn_text.draw_text(screen, choose_map_x, choose_map_y)
    exit_but.draw_but(screen)

    if lan:
        set1_but.draw_but(screen)
        set2_but.draw_but(screen)
        set3_but.draw_but(screen)


    else:
        set1vn_but.draw_but(screen)
        set2vn_but.draw_but(screen)
        set3vn_but.draw_but(screen)

    if set1_but.image_rect.collidepoint(pos):
        demo_map1_but.draw_but(screen)
        if pygame.mouse.get_pressed()[0] == 1:
            # Nv11_but.draw_but(screen)
            click = pygame.mixer.Sound("music/click.mp3")
            click.play(0)
            off_screen_except(12)
            choose_nv_screen1()

    if set2_but.image_rect.collidepoint(pos):
        demo_map2_but.draw_but(screen)
        if pygame.mouse.get_pressed()[0] == 1:
            click = pygame.mixer.Sound("music/click.mp3")
            click.play(0)
            off_screen_except(15)
            choose_nv_screen2()

    if set3_but.image_rect.collidepoint(pos):
        demo_map3_but.draw_but(screen)
        if pygame.mouse.get_pressed()[0] == 1:
            click = pygame.mixer.Sound("music/click.mp3")
            click.play(0)
            off_screen_except(21)
            choose_nv_screen3()

    if check_press(exit_but.image_rect, pos):
        off_screen_except(0)

    # Reset bien khi ve menu chon nhan vat
    global w11_x, w12_x, w13_x, w14_x, w15_x, wi11, wi12, wi13, wi14, wi15
    w11_x = 1
    w12_x = 1
    w13_x = 1
    w14_x = 1
    w15_x = 1
    wi11 = 0
    wi12 = 0
    wi13 = 0
    wi14 = 0
    wi15 = 0

    global random_x1, random_x2, random_x3, random_x4, random_x5, box_but1, box_but2, box_but3, box_but4, box_but5
    random_x1 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
    random_x2 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
    random_x3 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
    random_x4 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)
    random_x5 = random.randrange(SCREEN_WIDTH // 3, SCREEN_WIDTH // 2)

    box_but1 = Button(box, random_x1, 235, 50, 50)
    box_but2 = Button(box, random_x2, 350, 50, 50)
    box_but3 = Button(box, random_x3, 460, 50, 50)
    box_but4 = Button(box, random_x4, 575, 50, 50)
    box_but5 = Button(box, random_x5, 675, 50, 50)

    global x1, x2, x3, x4, x5
    x1 = random.uniform(0.8, 1.2)
    x2 = random.uniform(0.8, 1.2)
    x3 = random.uniform(0.8, 1.2)
    x4 = random.uniform(0.8, 1.2)
    x5 = random.uniform(0.8, 1.2)

    global lucky1, lucky2, lucky3, lucky4, lucky5
    lucky1 = random.randrange(1, 100)
    lucky2 = random.randrange(1, 100)
    lucky3 = random.randrange(1, 100)
    lucky4 = random.randrange(1, 100)
    lucky5 = random.randrange(1, 100)

    reset()


def setting_screen():
    background0.draw_bg(screen)
    exit_but.draw_but(screen)

    if lan:
        logout_but.draw_but(screen)
        language_but.draw_but(screen)
    else:
        logoutvn_but.draw_but(screen)
        languagevn_but.draw_but(screen)

    if check_press(logout_but.image_rect, pos):
        with open('log.txt', 'w') as file:
            file.write('0')
        off_screen_except(0)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0)
    if check_press(language_but.image_rect, pos):
        language_screen()


bet = ""

global nv
nv = 0
start_img = pygame.image.load("start/start.png")
start_but = Button(start_img, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, 200, 200)

more_info = pygame.image.load("Button/more_info.png")
more_info_but = Button(more_info, 1290, 325, 30, 30)


def choose_nv_screen1():
    global start_t
    global nv, setnv11, setnv21

    background0.draw_bg(screen)

    if lan:
        choosenv.draw_text(screen, 300, 10)
    else:
        choosenv_vn.draw_text(screen, 300, 10)
    if lan:
        set1_text.draw_text(screen, 50, 200)
        set2_text.draw_text(screen, 50, 400)
    else:
        set1vn_text.draw_text(screen, 50, 200)
        set2vn_text.draw_text(screen, 50, 400)

    # set 1
    if check_press(pygame.Rect(50, 200, 200, 80), pos):
        setnv11 = True
        setnv21 = False
    # set 2
    if check_press(pygame.Rect(50, 400, 200, 80), pos):
        setnv21 = True
        setnv11 = False

    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
        time.sleep(0.2)

    global s1_active, s2_active, s3_active, bet
    global info
    if not s1_active:
        global finish
        global background1
        Size1.draw_text(screen, 200, 600)
    else:
        background1 = background1
        finish = 1120
        size1_active.draw_text(screen, 200, 600)
    if not s2_active:
        Size2.draw_text(screen, 100, 600)

    else:
        background1 = background1_lv2
        finish = 1100
        size2_active.draw_text(screen, 100, 600)
    if not s3_active:
        Size3.draw_text(screen, 0, 600)
    else:
        finish = 1072
        background1 = background1_lv3
        size3_active.draw_text(screen, 0, 600)

    for event in pygame.event.get():
        if event.type == pygame.QUIT:
            pygame.quit()
            sys.exit()
        if event.type == pygame.KEYDOWN:
            if event.key == pygame.K_RETURN:

                # Khi nhấn Enter, in giá trị và đặt lại biến coin
                print("Coin value:", bet)
                bet = ""
            elif event.key == pygame.K_BACKSPACE:

                bet = bet[:-1]
            else:
                bet += event.unicode

    pygame.draw.rect(screen, (0, 0, 0), (600, 600, 300, 50), 2)

    text_surface = game_font1.render(bet, True, (255, 255, 255))
    screen.blit(text_surface, (600, 600))

    if setnv11:
        global x1, x2, x3, x4, x5
        global nv
        if check_press(Nv11_but.image_rect, pos):
            nv = 1
            if buff != 0:
                x1 = random.uniform(0.8, 1.2) * buff
            else:
                x1 = random.uniform(0.8, 1.2)
        if check_press(Nv12_but.image_rect, pos):
            nv = 2
            if buff != 0:
                x2 = random.uniform(0.8, 1.2) * buff
            else:
                x2 = random.uniform(0.8, 1.2)
        if check_press(Nv13_but.image_rect, pos):
            nv = 3
            if buff != 0:
                x3 = random.uniform(0.8, 1.2) * buff
            else:
                x3 = random.uniform(0.8, 1.2)
        if check_press(Nv14_but.image_rect, pos):
            nv = 4
            if buff != 0:
                x4 = random.uniform(0.8, 1.2) * buff
            else:
                x4 = random.uniform(0.8, 1.2)
        if check_press(Nv15_but.image_rect, pos):
            nv = 5
            if buff != 0:
                x5 = random.uniform(0.8, 1.2) * buff
            else:
                x5 = random.uniform(0.8, 1.2)
        if lan:
            nv1_choosed = Text(game_font1, "Character: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)
        else:
            nv1_choosed = Text(game_font1, "Nhân vật: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)

        Nv11_but.draw_but(screen)
        Nv12_but.draw_but(screen)
        Nv13_but.draw_but(screen)
        Nv14_but.draw_but(screen)
        Nv15_but.draw_but(screen)

        more_info_but.draw_but(screen)
        if more_info_but.image_rect.collidepoint(pos):
            if pygame.mouse.get_pressed()[0] == 1:
                nv_info(1, lan)

    if setnv21:
        if check_press(Nv11_but.image_rect, pos):
            nv = 1
            if buff != 0:
                x1 = random.uniform(0.8, 1.2) * buff
            else:
                x1 = random.uniform(0.8, 1.2)
        if check_press(Nv12_but.image_rect, pos):
            nv = 2
            if buff != 0:
                x2 = random.uniform(0.8, 1.2) * buff
            else:
                x2 = random.uniform(0.8, 1.2)
        if check_press(Nv13_but.image_rect, pos):
            nv = 3
            if buff != 0:
                x3 = random.uniform(0.8, 1.2) * buff
            else:
                x3 = random.uniform(0.8, 1.2)
        if check_press(Nv14_but.image_rect, pos):
            nv = 4
            if buff != 0:
                x4 = random.uniform(0.8, 1.2) * buff
            else:
                x4 = random.uniform(0.8, 1.2)
        if check_press(Nv15_but.image_rect, pos):
            nv = 5
            if buff != 0:
                x5 = random.uniform(0.8, 1.2) * buff
            else:
                x5 = random.uniform(0.8, 1.2)

        if lan:
            nv1_choosed = Text(game_font1, "Character: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)
        else:
            nv1_choosed = Text(game_font1, "Nhân vật: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)

        Nv41_but.draw_but(screen)
        Nv42_but.draw_but(screen)
        Nv43_but.draw_but(screen)
        Nv44_but.draw_but(screen)
        Nv45_but.draw_but(screen)

        more_info_but.draw_but(screen)
        if more_info_but.image_rect.collidepoint(pos):
            if pygame.mouse.get_pressed()[0] == 1:
                background0.draw_bg(screen)
                nv_info(4, lan)

    '''pygame.display.flip()'''
    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
    if pygame.rect.Rect(200, 600, 100, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s1_active = True
            s2_active = False
            s3_active = False
    if pygame.rect.Rect(100, 600, 80, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s2_active = True
            s1_active = False
            s3_active = False
    if pygame.rect.Rect(0, 600, 80, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s3_active = True
            s1_active = False
            s2_active = False
    '''pygame.display.update()'''

    next_but = Button(next_img, 1200, 720, 100, 50)
    if bet != "" and int(bet) <= int(coin) and nv != "" and 1 <= int(nv) <= 5:
        next_but.draw_but(screen)
        if check_press(next_but.image_rect, pos):
            global write_history_active
            write_history_active = True
            if setnv11:
                write_history_active = True
                global i
                i = 1
                box_active1 = box_active2 = box_active3 = box_active4 = box_active5 = 0
                global time_limit, start_time
                global three_limit, two_limit, one_limit
                three_limit = 1000
                two_limit = 2000
                one_limit = 3000
                time_limit = 4000
                start_time = pygame.time.get_ticks()
                off_screen_except(1)
                setnv11 = True
            if setnv21:
                write_history_active = True
                i = 1
                global box_active41, box_active42, box_active43, box_active44, box_active45
                box_active41 = box_active42 = box_active43 = box_active44 = box_active45 = 0
                three_limit = 1000
                two_limit = 2000
                one_limit = 3000
                time_limit = 4000
                start_time = pygame.time.get_ticks()
                off_screen_except(1)
                setnv21 = True


global nv2
nv2 = 0


def choose_nv_screen2():
    global setnv12, setnv22, nv2
    global s1_active, s2_active, s3_active, bet
    background0.draw_bg(screen)
    if lan:
        choosenv.draw_text(screen, 300, 10)
    else:
        choosenv_vn.draw_text(screen, 300, 10)
    if lan:
        set1_text.draw_text(screen, 50, 200)
        set2_text.draw_text(screen, 50, 400)
    else:
        set1vn_text.draw_text(screen, 50, 200)
        set2vn_text.draw_text(screen, 50, 400)
    if check_press(pygame.Rect(50, 200, 200, 80), pos):
        setnv12 = True
        setnv22 = False
    if check_press(pygame.Rect(50, 400, 200, 80), pos):
        setnv22 = True
        setnv12 = False
    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
        time.sleep(0.2)

    if not s1_active:
        global finish
        global background1
        Size1.draw_text(screen, 200, 600)
    else:
        background1 = background1
        finish = 1120
        size1_active.draw_text(screen, 200, 600)
    if not s2_active:
        Size2.draw_text(screen, 100, 600)

    else:
        background1 = background1_lv2
        finish = 1100
        size2_active.draw_text(screen, 100, 600)
    if not s3_active:
        Size3.draw_text(screen, 0, 600)
    else:
        finish = 1072
        background1 = background1_lv3
        size3_active.draw_text(screen, 0, 600)

    for event in pygame.event.get():
        if event.type == pygame.QUIT:
            pygame.quit()
            sys.exit()
        if event.type == pygame.KEYDOWN:
            if event.key == pygame.K_RETURN:

                # Khi nhấn Enter, in giá trị và đặt lại biến coin
                print("Coin value:", bet)
                bet = ""
            elif event.key == pygame.K_BACKSPACE:

                bet = bet[:-1]
            else:
                bet += event.unicode

    pygame.draw.rect(screen, (0, 0, 0), (600, 600, 300, 50), 2)
    exit_but.draw_but(screen)
    text_surface = game_font1.render(bet, True, (255, 255, 255))
    screen.blit(text_surface, (600, 600))

    if setnv12:
        global x1, x2, x3, x4, x5
        global nv2
        if check_press(Nv11_but.image_rect, pos):
            nv2 = 1
            if buff != 0:
                x1 = random.uniform(0.8, 1.2) * buff
            else:
                x1 = random.uniform(0.8, 1.2)
        if check_press(Nv12_but.image_rect, pos):
            nv2 = 2
            if buff != 0:
                x2 = random.uniform(0.8, 1.2) * buff
            else:
                x2 = random.uniform(0.8, 1.2)
        if check_press(Nv13_but.image_rect, pos):
            nv2 = 3
            if buff != 0:
                x3 = random.uniform(0.8, 1.2) * buff
            else:
                x3 = random.uniform(0.8, 1.2)
        if check_press(Nv14_but.image_rect, pos):
            nv2 = 4
            if buff != 0:
                x4 = random.uniform(0.8, 1.2) * buff
            else:
                x4 = random.uniform(0.8, 1.2)
        if check_press(Nv15_but.image_rect, pos):
            nv2 = 5
            if buff != 0:
                x5 = random.uniform(0.8, 1.2) * buff
            else:
                x5 = random.uniform(0.8, 1.2)

        if lan:
            nv1_choosed = Text(game_font1, "Character: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)
        else:
            nv1_choosed = Text(game_font1, "Nhân vật: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)

        Nv21_but.draw_but(screen)
        Nv22_but.draw_but(screen)
        Nv23_but.draw_but(screen)
        Nv24_but.draw_but(screen)
        Nv25_but.draw_but(screen)

        more_info_but.draw_but(screen)
        if more_info_but.image_rect.collidepoint(pos):
            if pygame.mouse.get_pressed()[0] == 1:
                background0.draw_bg(screen)
                nv_info(2, lan)

    if setnv22:
        if check_press(Nv11_but.image_rect, pos):
            nv2 = 1
            if buff != 0:
                x1 = random.uniform(0.8, 1.2) * buff
            else:
                x1 = random.uniform(0.8, 1.2)
        if check_press(Nv12_but.image_rect, pos):
            nv2 = 2
            if buff != 0:
                x2 = random.uniform(0.8, 1.2) * buff
            else:
                x2 = random.uniform(0.8, 1.2)
        if check_press(Nv13_but.image_rect, pos):
            nv2 = 3
            if buff != 0:
                x3 = random.uniform(0.8, 1.2) * buff
            else:
                x3 = random.uniform(0.8, 1.2)
        if check_press(Nv14_but.image_rect, pos):
            nv2 = 4
            if buff != 0:
                x4 = random.uniform(0.8, 1.2) * buff
            else:
                x4 = random.uniform(0.8, 1.2)
        if check_press(Nv15_but.image_rect, pos):
            nv2 = 5
            if buff != 0:
                x5 = random.uniform(0.8, 1.2) * buff
            else:
                x5 = random.uniform(0.8, 1.2)
        if lan:
            nv1_choosed = Text(game_font1, "Character: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)
        else:
            nv1_choosed = Text(game_font1, "Nhân vật: " + str(nv), 0, 0, 0)
            nv1_choosed.draw_text(screen, 500, 500)

        Nv51_but.draw_but(screen)
        Nv52_but.draw_but(screen)
        Nv53_but.draw_but(screen)
        Nv54_but.draw_but(screen)
        Nv55_but.draw_but(screen)

        more_info_but.draw_but(screen)
        if more_info_but.image_rect.collidepoint(pos):
            if pygame.mouse.get_pressed()[0] == 1:
                background0.draw_bg(screen)
                nv_info(5, lan)

    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
    if pygame.rect.Rect(200, 600, 100, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s1_active = True
            s2_active = False
            s3_active = False
    if pygame.rect.Rect(100, 600, 80, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s2_active = True
            s1_active = False
            s3_active = False
    if pygame.rect.Rect(0, 600, 80, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s3_active = True
            s1_active = False
            s2_active = False
    '''pygame.display.update()'''
    next_img = pygame.image.load("Button/next.png")
    next_but = Button(next_img, 1200, 720, 100, 50)
    if bet != "" and int(bet) <= coin and nv2 != "" and 1 <= int(nv2) <= 5:
        next_but.draw_but(screen)
        next_but.draw_but(screen)
        if check_press(next_but.image_rect, pos):
            if setnv12:
                global box_active21, box_active22, box_active23, box_active24, box_active25
                box_active21 = box_active22 = box_active23 = box_active24 = box_active25 = 0
                global time_limit, start_time
                global three_limit, two_limit, one_limit
                three_limit = 1000
                two_limit = 2000
                one_limit = 3000
                time_limit = 4000
                start_time = pygame.time.get_ticks()
                off_screen_except(2)
                setnv12 = True
            if setnv22:
                global box_active51, box_active52, box_active53, box_active54, box_active55
                box_active51 = box_active52 = box_active53 = box_active54 = box_active55 = 0
                three_limit = 1000
                two_limit = 2000
                one_limit = 3000
                time_limit = 4000
                start_time = pygame.time.get_ticks()
                off_screen_except(2)
                setnv22 = True


nv3 = 0


def choose_nv_screen3():
    global s1_active, s2_active, s3_active, bet, nv3

    background0.draw_bg(screen)
    if lan:
        choosenv.draw_text(screen, 300, 10)
    else:
        choosenv_vn.draw_text(screen, 300, 10)
    Nv31_but.draw_but(screen)
    Nv32_but.draw_but(screen)
    Nv33_but.draw_but(screen)
    Nv34_but.draw_but(screen)
    Nv35_but.draw_but(screen)
    if check_press(Nv31_but.image_rect, pos):
        nv3 = 1
    if check_press(Nv32_but.image_rect, pos):
        nv3 = 2
    if check_press(Nv33_but.image_rect, pos):
        nv3 = 3
    if check_press(Nv34_but.image_rect, pos):
        nv3 = 4
    if check_press(Nv35_but.image_rect, pos):
        nv3 = 5
    if lan:
        nv1_choosed = Text(game_font1, "Character: " + str(nv3), 0, 0, 0)
        nv1_choosed.draw_text(screen, 500, 500)
    else:
        nv1_choosed = Text(game_font1, "Nhân vật: " + str(nv3), 0, 0, 0)
        nv1_choosed.draw_text(screen, 500, 500)
    if pygame.rect.Rect(200, 600, 100, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s1_active = True
            s2_active = False
            s3_active = False
    if pygame.rect.Rect(100, 600, 80, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s2_active = True
            s1_active = False
            s3_active = False
    if pygame.rect.Rect(0, 600, 80, 40).collidepoint(pos):
        if pygame.mouse.get_pressed()[0] == 1:
            s3_active = True
            s1_active = False
            s2_active = False

    if not s1_active:
        global finish
        global background1
        Size1.draw_text(screen, 200, 600)
    else:
        background1 = background1
        finish = 1120
        size1_active.draw_text(screen, 200, 600)
    if not s2_active:
        Size2.draw_text(screen, 100, 600)

    else:
        background1 = background1_lv2
        finish = 1100
        size2_active.draw_text(screen, 100, 600)
    if not s3_active:
        Size3.draw_text(screen, 0, 600)
    else:
        finish = 1072
        background1 = background1_lv3
        size3_active.draw_text(screen, 0, 600)
    for event in pygame.event.get():
        if event.type == pygame.QUIT:
            pygame.quit()
            sys.exit()
        if event.type == pygame.KEYDOWN:
            if event.key == pygame.K_RETURN:

                # Khi nhấn Enter, in giá trị và đặt lại biến coin
                print("Coin value:", bet)
                bet = ""
            elif event.key == pygame.K_BACKSPACE:

                bet = bet[:-1]
            else:
                bet += event.unicode
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)

    '''pygame.display.update()'''
    next_img = pygame.image.load("Button/next.png")
    next_but = Button(next_img, 1200, 720, 100, 50)
    if bet != "" and int(bet) <= coin and nv3 != "" and 1 <= int(nv3) <= 5:
        next_but.draw_but(screen)
        next_but.draw_but(screen)
        if check_press(next_but.image_rect, pos):
            global time_limit, start_time
            global three_limit, two_limit, one_limit
            three_limit = 1000
            two_limit = 2000
            one_limit = 3000
            time_limit = 4000
            start_time = pygame.time.get_ticks()
            off_screen_except(3)

    pygame.draw.rect(screen, (0, 0, 0), (600, 600, 300, 50), 2)

    text_surface = game_font1.render(bet, True, (255, 255, 255))
    screen.blit(text_surface, (600, 600))
    pygame.display.flip()

    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
        time.sleep(0.2)


s1_active = True
s2_active = s3_active = False

nv1_choosed = nv2_choosed = nv3_choosed = nv4_choosed = nv5_choosed = False

box_but1 = Button(box, random_x1, 235, 30, 30)
box_but2 = Button(box, random_x2, 350, 50, 50)
box_but3 = Button(box, random_x3, 460, 50, 50)
box_but4 = Button(box, random_x4, 575, 50, 50)
box_but5 = Button(box, random_x5, 675, 50, 50)


def nv_info(x):
    if x == 1:
        set1_infovn = pygame.image.load("story/11.png")
        set1_info = pygame.image.load("story/1.png")
        bg = Background(set1_info)
        bg1 = Background(set1_infovn)
        if lan:
            bg.draw_bg(screen)
        else:
            bg1.draw_bg(screen)

    if x == 2:
        set2_infovn = pygame.image.load("story/9.png")
        set2_info = pygame.image.load("story/10.png")
        bg = Background(set2_info)
        bg1 = Background(set2_infovn)
        if lan:
            bg.draw_bg(screen)
        else:
            bg1.draw_bg(screen)
    if x == 4:
        set4_infovn = pygame.image.load("story/13.png")
        set4_info = pygame.image.load("story/14.png")
        bg = Background(set4_info)
        bg1 = Background(set4_infovn)
        if lan:
            bg.draw_bg(screen)
        else:
            bg1.draw_bg(screen)
    if x == 5:
        set5_infovn = pygame.image.load("story/7.png")
        set5_info = pygame.image.load("story/8.png")
        bg = Background(set5_info)
        bg1 = Background(set5_infovn)
        if lan:
            bg.draw_bg(screen)
        else:
            bg1.draw_bg(screen)


def map1(pos):
    background1.draw_bg(screen)
    board_but.draw_but(screen)
    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
        time.sleep(0.2)

    S_nv1.draw_text(screen, 1230, 20)
    S_nv2.draw_text(screen, 1230, 45)
    S_nv3.draw_text(screen, 1230, 70)
    S_nv4.draw_text(screen, 1230, 95)
    S_nv5.draw_text(screen, 1230, 120)

    if setnv11:
        head11_but = Button(head11, 1180, 35, 26, 26)
        head12_but = Button(head12, 1180, 60, 29, 29)
        head13_but = Button(head13, 1180, 85, 25, 25)
        head14_but = Button(head14, 1180, 110, 24, 22)
        head15_but = Button(head15, 1180, 135, 24, 22)

        head11_but.draw_but(screen)
        head12_but.draw_but(screen)
        head13_but.draw_but(screen)
        head14_but.draw_but(screen)
        head15_but.draw_but(screen)
    if setnv21:
        head41_but = Button(head41, 1180, 35, 26, 26)
        head42_but = Button(head42, 1180, 60, 29, 29)
        head43_but = Button(head43, 1180, 85, 25, 25)
        head44_but = Button(head44, 1180, 110, 24, 22)
        head45_but = Button(head45, 1180, 135, 24, 22)

        head41_but.draw_but(screen)
        head42_but.draw_but(screen)
        head43_but.draw_but(screen)
        head44_but.draw_but(screen)
        head45_but.draw_but(screen)

    box_but1 = Button(box, random_x1, 235, 50, 50)
    box_but2 = Button(box, random_x2, 350, 50, 50)
    box_but3 = Button(box, random_x3, 460, 50, 50)
    box_but4 = Button(box, random_x4, 575, 50, 50)
    box_but5 = Button(box, random_x5, 675, 50, 50)

    if box_touch[1] == 0:
        box_but1.draw_but(screen)
    if box_touch[2] == 0:
        box_but2.draw_but(screen)
    if box_touch[3] == 0:
        box_but3.draw_but(screen)
    if box_touch[4] == 0:
        box_but4.draw_but(screen)
    if box_touch[5] == 0:
        box_but5.draw_but(screen)

    cur_time = pygame.time.get_ticks()
    if cur_time - start_t < 2000:
        win_but.draw_but(screen)


def map2(pos):
    background2.draw_bg(screen)
    board_but.draw_but(screen)
    exit_but.draw_but(screen)
    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
        time.sleep(0.2)
    S_nv1.draw_text(screen, 1230, 20)
    S_nv2.draw_text(screen, 1230, 45)
    S_nv3.draw_text(screen, 1230, 70)
    S_nv4.draw_text(screen, 1230, 95)
    S_nv5.draw_text(screen, 1230, 120)

    box_but1 = Button(box, random_x1, 235, 50, 50)
    box_but2 = Button(box, random_x2, 350, 50, 50)
    box_but3 = Button(box, random_x3, 460, 50, 50)
    box_but4 = Button(box, random_x4, 575, 50, 50)
    box_but5 = Button(box, random_x5, 675, 50, 50)

    if setnv12:
        head31_but = Button(head31, 1180, 35, 26, 26)
        head32_but = Button(head32, 1180, 60, 29, 29)
        head33_but = Button(head33, 1180, 85, 25, 25)
        head34_but = Button(head34, 1180, 110, 24, 22)
        head35_but = Button(head35, 1180, 135, 24, 22)

        head31_but.draw_but(screen)
        head32_but.draw_but(screen)
        head33_but.draw_but(screen)
        head34_but.draw_but(screen)
        head35_but.draw_but(screen)

    if setnv22:
        head51_but = Button(head51, 1180, 35, 26, 26)
        head52_but = Button(head52, 1180, 60, 29, 29)
        head53_but = Button(head53, 1180, 85, 25, 25)
        head54_but = Button(head54, 1180, 110, 24, 22)
        head55_but = Button(head55, 1180, 135, 24, 22)

        head51_but.draw_but(screen)
        head52_but.draw_but(screen)
        head53_but.draw_but(screen)
        head54_but.draw_but(screen)
        head55_but.draw_but(screen)

    if box_touch[1] == 0:
        box_but1.draw_but(screen)
    if box_touch[2] == 0:
        box_but2.draw_but(screen)
    if box_touch[3] == 0:
        box_but3.draw_but(screen)
    if box_touch[4] == 0:
        box_but4.draw_but(screen)
    if box_touch[5] == 0:
        box_but5.draw_but(screen)


def map3(pos):
    off_screen_except(3)
    background3.draw_bg(screen)
    board_but.draw_but(screen)
    exit_but.draw_but(screen)

    if check_press(exit_but.image_rect, pos):
        off_screen_except(0.5)
        time.sleep(0.2)

    S_nv1.draw_text(screen, 1230, 20)
    S_nv2.draw_text(screen, 1230, 45)
    S_nv3.draw_text(screen, 1230, 70)
    S_nv4.draw_text(screen, 1230, 95)
    S_nv5.draw_text(screen, 1230, 120)

    box_but1 = Button(box, random_x1, 235, 50, 50)
    box_but2 = Button(box, random_x2, 350, 50, 50)
    box_but3 = Button(box, random_x3, 460, 50, 50)
    box_but4 = Button(box, random_x4, 575, 50, 50)
    box_but5 = Button(box, random_x5, 675, 50, 50)

    if box_touch[1] == 0:
        box_but1.draw_but(screen)
    if box_touch[2] == 0:
        box_but2.draw_but(screen)
    if box_touch[3] == 0:
        box_but3.draw_but(screen)
    if box_touch[4] == 0:
        box_but4.draw_but(screen)
    if box_touch[5] == 0:
        box_but5.draw_but(screen)
    head21_but = Button(head21, 1180, 35, 26, 26)
    head22_but = Button(head22, 1180, 60, 29, 29)
    head23_but = Button(head23, 1180, 85, 25, 25)
    head24_but = Button(head24, 1180, 110, 24, 22)
    head25_but = Button(head25, 1180, 135, 24, 22)

    head21_but.draw_but(screen)
    head22_but.draw_but(screen)
    head23_but.draw_but(screen)
    head24_but.draw_but(screen)
    head25_but.draw_but(screen)


clock = pygame.time.Clock()
three_limit = 1000  # milliseconds (5 seconds)
start_time = pygame.time.get_ticks()
two_limit = 2000
one_limit = 3000
time_limit = 4000

start_img = pygame.image.load("start/start.png")
start_but = Button(start_img, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, 350, 250)

startvn_img = pygame.image.load("start/startvn.png")
start_butvn = Button(startvn_img, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, 350, 250)
three_img = pygame.image.load("start/3.png")
two_img = pygame.image.load("start/2.png")
one_img = pygame.image.load("start/1.png")

three_but = Button(three_img, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, 200, 200)
two_but = Button(two_img, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, 200, 200)
one_but = Button(one_img, SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2, 200, 200)

start_but.draw_but(screen)
three_but.draw_but(screen)
two_but.draw_but(screen)
one_but.draw_but(screen)
list_nv_win = [0, 0, 0, 0, 0]
list_ranking = []
list_stt = []
shop = store_func.Store(lan, 1344, 768)
get_coin = False

save = pygame.image.load("Button/save_ but.png")
save_but = Button(save, 1300, 720, 80, 80)
from docx import Document


def clear_word_file(file_path):
    # Mở tệp Word
    doc = Document(file_path)

    # Xóa tất cả các đoạn văn bản trong tài liệu
    for paragraph in doc.paragraphs:
        paragraph.clear()

    # Xóa tất cả các bảng trong tài liệu
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.clear()

    # Lưu lại tệp Word đã xóa sạch
    doc.save(file_path)


onevn = pygame.image.load("NPC/Vie/20.png")
twovn = pygame.image.load("NPC/Vie/21.png")
threevn = pygame.image.load("NPC/Vie/23.png")
fourvn = pygame.image.load("NPC/Vie/25.png")
fivevn = pygame.image.load("NPC/Vie/27.png")
sixvn = pygame.image.load("NPC/Vie/29.png")
A = [onevn, twovn, threevn, fourvn, fivevn, sixvn]
for i in range(0, 6, 1):
    A[i] = Background(A[i])

one = pygame.image.load("NPC/Eng/19.png")
two = pygame.image.load("NPC/Eng/22.png")
three = pygame.image.load("NPC/Eng/24.png")
four = pygame.image.load("NPC/Eng/26.png")
five = pygame.image.load("NPC/Eng/28.png")
six = pygame.image.load("NPC/Eng/30.png")
B = [one, two, three, four, five, six]
for j in range(0, 6, 1):
    B[j] = Background(B[j])


def first_login_screen(lan, i, screen):
    if lan:
        B[i].draw_bg(screen)
    else:
        A[i].draw_bg(screen)


h_y = 200
head_array = [0, 0, 0, 0, 0]


def ketqua_screen(map, set, array_stt):
    global h_y
    if lan:
        result_screen = pygame.image.load("16.png")
        result_screen_bg = Background(result_screen)
        result_screen_bg.draw_bg(screen)
    else:
        result_screen = pygame.image.load("15.png")
        result_screen_bg = Background(result_screen)
        result_screen_bg.draw_bg(screen)
    '''pygame.image.save(screen, "temp.png")'''  # Chụp trước khi vẽ nút save


    if check_press(exit_but.image_rect, pos):
        global list_stt
        list_stt = []
        pygame.mixer.music.stop()
        pygame.mixer.music.load("music/nhacnengame.mp3")
        pygame.mixer.music.play(-1)
        background = Background(bg1)
        off_screen_except(0)

    if map == 1 and set == 1:
        h1 = Button(head11, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head12, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head13, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head14, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head15, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        global list_nv_win, list_ranking
        h_y = 200
        h1 = Button(head11, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head12, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head13, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head14, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head15, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[0]) - 1].draw_but(screen)
        h_y = 380
        h1 = Button(head11, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head12, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head13, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head14, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head15, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[1]) - 1].draw_but(screen)
        h_y = 580
        h1 = Button(head11, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head12, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head13, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head14, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head15, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[2]) - 1].draw_but(screen)

    if map == 1 and set == 2:
        h1 = Button(head41, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head42, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head43, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head44, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head45, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        global list_nv_win, list_ranking
        h_y = 200
        h1 = Button(head41, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head42, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head43, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head44, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head45, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[0]) - 1].draw_but(screen)
        h_y = 380
        h1 = Button(head41, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head42, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head43, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head44, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head45, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[1]) - 1].draw_but(screen)
        h_y = 580
        h1 = Button(head41, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head42, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head43, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head44, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head45, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[2]) - 1].draw_but(screen)

    if map == 2 and set == 1:
        h1 = Button(head31, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head32, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head33, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head34, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head35, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        global list_nv_win, list_ranking
        h_y = 200
        h1 = Button(head31, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head32, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head33, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head34, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head35, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[0]) - 1].draw_but(screen)
        h_y = 380
        h1 = Button(head31, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head32, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head33, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head34, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head35, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[1]) - 1].draw_but(screen)
        h_y = 580
        h1 = Button(head31, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head32, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head33, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head34, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head35, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[2]) - 1].draw_but(screen)

    if map == 2 and set == 2:
        h1 = Button(head51, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head52, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head53, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head54, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head55, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        global list_nv_win, list_ranking
        h_y = 200
        h1 = Button(head51, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head52, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head53, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head54, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head55, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[0]) - 1].draw_but(screen)
        h_y = 380
        h1 = Button(head51, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head52, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head53, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head54, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head55, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[1]) - 1].draw_but(screen)
        h_y = 580
        h1 = Button(head51, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head52, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head53, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head54, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head55, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[2]) - 1].draw_but(screen)

    if map == 3:
        h1 = Button(head21, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head22, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head23, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head24, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head25, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        global list_nv_win, list_ranking
        h_y = 200
        h1 = Button(head21, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head22, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head23, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head24, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head25, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[0]) - 1].draw_but(screen)
        h_y = 380
        h1 = Button(head21, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head22, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head23, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head24, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head25, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[1]) - 1].draw_but(screen)
        h_y = 580
        h1 = Button(head21, 1344 // 2.5, h_y, 80, 80)
        h2 = Button(head22, 1344 // 2.5, h_y, 80, 80)
        h3 = Button(head23, 1344 // 2.5, h_y, 80, 80)
        h4 = Button(head24, 1344 // 2.5, h_y, 80, 80)
        h5 = Button(head25, 1344 // 2.5, h_y, 80, 80)
        head_array = [h1, h2, h3, h4, h5]
        head_array[int(array_stt[2]) - 1].draw_but(screen)
    if check_press(save_but.image_rect, pos):
        screenshot.capture_screen(screen, user)
        time.sleep(0.5)
    save_but.draw_but(screen)
    exit_but.draw_but(screen)

    list_nv_win = [0, 0, 0, 0, 0]
    list_ranking = []

    '''list_stt = []'''


set1_info = pygame.image.load("nvinfo/info1.png")
set2_info = pygame.image.load("nvinfo/info2.png")
set4_info = pygame.image.load("nvinfo/info4.png")
set5_info = pygame.image.load("nvinfo/info5.png")

set1_info_bg = Background(set1_info)
set2_info_bg = Background(set2_info)
set4_info_bg = Background(set4_info)
set5_info_bg = Background(set5_info)

set1_infovn = pygame.image.load("nvinfo/info1vn.png")
set2_infovn = pygame.image.load("nvinfo/info2vn.png")
set4_infovn = pygame.image.load("nvinfo/info4vn.png")
set5_infovn = pygame.image.load("nvinfo/info5vn.png")

set1_infovn_bg = Background(set1_infovn)
set2_infovn_bg = Background(set2_infovn)
set4_infovn_bg = Background(set4_infovn)
set5_infovn_bg = Background(set5_infovn)


# Màn hình info nhân vật
def nv_info(set, lan):
    if lan:
        if set == 1: set1_info_bg.draw_bg(screen)
        if set == 2: set2_info_bg.draw_bg(screen)
        if set == 4: set4_info_bg.draw_bg(screen)
        if set == 5: set5_info_bg.draw_bg(screen)
    else:
        if set == 1: set1_infovn_bg.draw_bg(screen)
        if set == 2: set2_infovn_bg.draw_bg(screen)
        if set == 4: set4_infovn_bg.draw_bg(screen)
        if set == 5: set5_infovn_bg.draw_bg(screen)


teleport = pygame.image.load("hieuung/teleport.png")
teleport = pygame.transform.scale(teleport, (100, 150))
stun = pygame.image.load("hieuung/stun.png")
stun = pygame.transform.scale(stun, (75, 50))
quicken = pygame.image.load("hieuung/quicken.png")
quicken = pygame.transform.scale(quicken, (50, 50))
boom = pygame.image.load("hieuung/boom.png")
boom = pygame.transform.scale(boom, (50, 50))
reverse = pygame.image.load("hieuung/reverse.png")
reverse = pygame.transform.scale(reverse, (50, 50))


def reset():  # hàm làm mới các hiệu ứng trước khi bắt đầu 1 game mới
    global quicken_atr, box_touch, rever, k_stun, tele, tele_x
    quicken_atr = [0, 0, 0, 0, 0, 0]
    box_touch = [0, 0, 0, 0, 0, 0]
    rever = [0, 0, 0, 0, 0, 0]
    k_stun = [0, 0, 0, 0, 0, 0]
    tele = [0, 0, 0, 0, 0, 0]
    tele_x = [0, 0, 0, 0, 0, 0]


def stun_effect(obj_x, m, frame, y):  # hiệu ứng làm choáng
    obj_x -= m
    frame -= 1
    t = False
    if (frame // 5) % 2: t = True
    if frame > 130:
        screen.blit(boom, (obj_x - 25, y - 25))
    else:
        screen.blit(pygame.transform.flip(stun, t, False), (obj_x - 25, y - 80))
    return obj_x, frame


def tele_effect(frame, x, dist, y):  # hiệu ứng teleport
    screen.blit(teleport, (x - 25, y - 75))
    screen.blit(teleport, (x + dist - 25, y - 75))
    return frame - 1


def lucky_atr(obj_x, dist, lucky, nv, i):  # làm gọn phần xử lí chỉ số may mắn
    box_touch[i] = 1
    if 1 <= lucky <= 30:  # tỉ lệ lần lượt là tăng tốc 30%
        dist = nv.buff_speed(0)  # choáng   25%
        quicken_atr[i] = 1  # quay đầu 25%
    if 31 <= lucky <= 55:  # tele     20%
        k_stun[i] = 150  # về đích   1%
    if 56 <= lucky <= 80:  # các mảng k_stun, rever, tele, quicken_atr là để đếm thời gian vẽ các hiệu ứng
        rever[i] = 100
    if 81 <= lucky <= 99:
        tele[i] = 20
        tele_x[i] = obj_x
        obj_x += lucky * 3
    if lucky == 100:
        tele[i] = 20
        tele_x[i] = obj_x
        lucky = (finish + 10 - obj_x) // 3
        obj_x = finish
    return obj_x, dist


def play_sound_effect(sound_file):
    sound = pygame.mixer.Sound(sound_file)
    sound.play()
    return sound


def play_background_music(music_file):
    pygame.mixer.music.load(music_file)
    pygame.mixer.music.play(-1)


start_t = 0
i = 0
play_background_music("music/nhacnengame.mp3")
reset()
k_stun = 0
music = 0
win_or_lose = False


def first_login_screen(lan, i, screen):
    if lan:
        B[i].draw_bg(screen)
    else:
        A[i].draw_bg(screen)


scroll_y = 0
angle = 0
banhxe1 = pygame.image.load("setbg3/set03/banhxe1.png")
banhxe2 = pygame.image.load("setbg3/set03/banhxe1.png")
banhxe3 = pygame.image.load("setbg3/set03/banhxe1.png")
banhxe5 = pygame.image.load("setbg3/set03/banhxe1.png")


def rotate(image_path, x, y, w, h):
    global angle
    image = pygame.image.load(image_path)
    image = pygame.transform.scale(image, (w, h))
    # Xoay hình ảnh
    rotated_image = pygame.transform.rotate(image, -angle)
    rotated_rect = rotated_image.get_rect(center=(x, y))

    # Vẽ hình ảnh đã xoay lên màn hình
    screen.blit(rotated_image, rotated_rect.topleft)

    angle += 1  # Tăng góc xoay mỗi lần vòng lặp
    if angle >= 360:
        angle = 0


while True:

    with open('user.txt', 'r') as file:
        user = file.read()
    with open("account/" + str(user) + '/coin.txt', 'r') as file:
        coin = int(file.read())
    with open('first_log.txt', 'r') as file:
        if str(file.read()) == "1":
            log_1st = True

    global pos
    pos = pygame.mouse.get_pos()

    S_nv1 = Text(game_font2, str(int(w11_x / 5)), 0, 0, 0)
    S_nv2 = Text(game_font2, str(int(w12_x / 5)), 0, 0, 0, )
    S_nv3 = Text(game_font2, str(int(w13_x / 5)), 0, 0, 0)
    S_nv4 = Text(game_font2, str(int(w14_x / 5)), 0, 0, 0)
    S_nv5 = Text(game_font2, str(int(w15_x / 5)), 0, 0, 0)

    if w11_x < finish + 30:
        w11_x += x1

    if w12_x < finish + 30:
        w12_x += x2

    if w13_x < finish + 30:
        w13_x += x3

    if w14_x < finish + 30:
        w14_x += x4

    if w15_x < finish + 30:
        w15_x += x5

    wi11 += 0.5
    wi12 += 0.5
    wi13 += 0.5
    wi14 += 0.5
    wi15 += 0.125
    if wi11 == 17: wi11 = 0
    if wi12 == 11: wi12 = 0
    if wi13 == 17: wi13 = 0
    if wi14 == 17: wi14 = 0
    if wi15 == 17: wi15 = 0

    with open('log.txt', 'r') as file:
        content = file.read()
        if content == "1":
            lock = True
        if content == "0":
            lock = False

    with open('lan.txt', 'r') as file:
        l = file.read()
        if l == "1":
            lan = True
        if l == "0":
            lan = False

    for ev in pygame.event.get():
        if ev.type == pygame.QUIT:
            pygame.quit()
            sys.exit()
        if ev.type == pygame.MOUSEBUTTONDOWN and history_bool:
            if ev.button == 4:
                scroll_y += 20
        if ev.type == pygame.MOUSEBUTTONDOWN and history_bool:
            if ev.button == 5:
                scroll_y -= 20

    if set_:
        with open('first_log.txt', 'r') as file:
            if file.read() == "1":
                log1st = True
        if not log1st:
            menu_screen()
            music = 0
            store = store_func
            shop = store.Store(lan, 1344, 756)
            store_img = pygame.image.load("extra_stuff/store/store.png")
            store = Button(store_img, 1200, 650, 200, 200)
            if lock:
                store.draw_but(screen)
            if check_press(store.image_rect, pos) and lock:
                pygame.mixer.music.stop()
                pygame.mixer.music.load("music/nhac nen mua do.mp3")
                pygame.mixer.music.play(-1)
                if game_finish == 1:
                    shop.buy1 = True
                    shop.buy2 = True
                    shop.buy3 = True
                    game_finish = 0
                buff, coin = shop.running(screen, coin, 1344, 768)
                coin_but.draw_but(screen)

                coin_text = Text(game_font1, str(coin), 255, 200, 60)
                coin_text.draw_text(screen, 1130, -6)
        else:
            first_login_screen(lan, i, screen)
            if check_press(pygame.rect.Rect(0, 0, 1344, 756), pos):
                i += 1
                if i == 6:
                    log1st = False
                    with open('first_log.txt', 'w') as file:
                        file.write("0")

    if wait:
        wait_screen()

    if setting_bool:
        setting_screen()
    if set2:
        music = 1
        current_time = pygame.time.get_ticks()
        elapsed_time = current_time - start_time

        if elapsed_time < 1000:
            background2.draw_bg(screen)
            pygame.mixer.music.stop()
            pygame.mixer.music.load("music/countdown1.mp3")
            pygame.mixer.music.play(0)
        elif elapsed_time < three_limit + 1000:
            background2.draw_bg(screen)
            three_but.draw_but(screen)
        elif elapsed_time < two_limit + 1000:
            background2.draw_bg(screen)
            two_but.draw_but(screen)
        elif elapsed_time < one_limit + 1000:
            background2.draw_bg(screen)
            one_but.draw_but(screen)
        elif elapsed_time < time_limit + 1000:
            background2.draw_bg(screen)

            if lan:
                start_but.draw_but(screen)
            else:
                start_butvn.draw_but(screen)
            pygame.mixer.music.stop()
            pygame.mixer.music.load("music/lucvodua2.mp3")
            pygame.mixer.music.play(-1)
            w11_x = w12_x = w13_x = w14_x = w15_x = 10

        else:

            boom_but1 = Button(boom, box_but1.x, 225, 90, 90)
            boom_but2 = Button(boom, box_but2.x, 345, 90, 90)
            boom_but3 = Button(boom, box_but3.x, 470, 90, 90)
            boom_but4 = Button(boom, box_but4.x, 570, 90, 90)
            boom_but5 = Button(boom, box_but5.x, 660, 90, 90)

            map2(pos)
            
            if check_press(exit_but.image_rect, pos):
                pygame.mixer.music.stop()
                pygame.mixer.music.load("music/nhacnengame.mp3")
                pygame.mixer.music.play(-1)
                off_screen_except(0.5)
                time.sleep(0.2)
            if setnv12:
                choosen_set = 2
                nv21 = nhanVat("nv21", w21_name[int(wi12)], w11_x, 200, 120, 120)
                nv21.draw_nv(1)
                if nv21.check_vc(box_but1) and box_touch[1] == 0: w11_x, x1 = lucky_atr(w11_x, x1, lucky1, nv21, 1)
                if k_stun[1]: w11_x, k_stun[1] = stun_effect(w11_x, x1, k_stun[1], 200)
                if rever[1]: w11_x -= 2 * x1
                if tele[1]: tele[1] = tele_effect(tele[1], tele_x[1], lucky1 * 3, 200)

                nv22 = nhanVat("nv22", w22_name[int(wi12)], w12_x, 325, 120, 120)
                nv22.draw_nv(2)
                if nv22.check_vc(box_but2) and box_touch[2] == 0: w12_x, x2 = lucky_atr(w12_x, x2, lucky2, nv22, 2)
                if k_stun[2]: w12_x, k_stun[2] = stun_effect(w12_x, x2, k_stun[2], 325)
                if rever[2]: w12_x -= 2 * x2
                if tele[2]: tele[2] = tele_effect(tele[2], tele_x[2], lucky2 * 3, 325)

                nv23 = nhanVat("nv23", w23_name[int(wi12)], w13_x, 450, 120, 120)
                nv23.draw_nv(3)
                if nv23.check_vc(box_but3) and box_touch[3] == 0: w13_x, x3 = lucky_atr(w13_x, x3, lucky3, nv23, 3)
                if k_stun[3]: w13_x, k_stun[3] = stun_effect(w13_x, x3, k_stun[3], 450)
                if rever[3]: w13_x -= 2 * x3
                if tele[3]: tele[3] = tele_effect(tele[3], tele_x[3], lucky3 * 3, 450)

                nv24 = nhanVat("nv24", w24_name[int(wi12)], w14_x, 550, 120, 120)
                nv24.draw_nv(4)
                if nv24.check_vc(box_but4) and box_touch[4] == 0: w14_x, x4 = lucky_atr(w14_x, x4, lucky4, nv24, 4)
                if k_stun[4]: w14_x, k_stun[4] = stun_effect(w14_x, x4, k_stun[4], 550)
                if rever[4]: w14_x -= 2 * x4
                if tele[4]: tele[4] = tele_effect(tele[4], tele_x[4], lucky4 * 3, 550)

                nv25 = nhanVat("nv25", w25_name[int(wi12)], w15_x, 650, 120, 120)
                nv25.draw_nv(5)
                if nv25.check_vc(box_but5) and box_touch[5] == 0: w15_x, x5 = lucky_atr(w15_x, x5, lucky5, nv25, 5)
                if k_stun[5]: w15_x, k_stun[5] = stun_effect(w15_x, x5, k_stun[5], 650)
                if rever[5]: w15_x -= 2 * x5
                if tele[5]: tele[5] = tele_effect(tele[5], tele_x[5], lucky5 * 3, 650)

            if setnv22:
                choosen_set = 5
                nv51 = nhanVat("nv51", w51_name[int(wi15)], w11_x, 200, 100, 100)
                nv51.draw_nv(1)
                if nv51.check_vc(box_but1) and box_touch[1] == 0: w11_x, x1 = lucky_atr(w11_x, x1, lucky1, nv51, 1)
                if k_stun[1]: w11_x, k_stun[1] = stun_effect(w11_x, x1, k_stun[1], 200)
                if rever[1]: w11_x -= 2 * x1
                if tele[1]: tele[1] = tele_effect(tele[1], tele_x[1], lucky1 * 3, 200)

                nv52 = nhanVat("nv52", w52_name[int(wi15)], w12_x, 325, 100, 100)
                nv52.draw_nv(2)
                if nv52.check_vc(box_but2) and box_touch[2] == 0: w12_x, x2 = lucky_atr(w12_x, x2, lucky2, nv52, 2)
                if k_stun[2]: w12_x, k_stun[2] = stun_effect(w12_x, x2, k_stun[2], 325)
                if rever[2]: w12_x -= 2 * x2
                if tele[2]: tele[2] = tele_effect(tele[2], tele_x[2], lucky2 * 3, 325)

                nv53 = nhanVat("nv53", w53_name[int(wi15)], w13_x, 450, 100, 100)
                nv53.draw_nv(3)
                if nv53.check_vc(box_but3) and box_touch[3] == 0: w13_x, x3 = lucky_atr(w13_x, x3, lucky3, nv53, 3)
                if k_stun[3]: w13_x, k_stun[3] = stun_effect(w13_x, x3, k_stun[3], 450)
                if rever[3]: w13_x -= 2 * x3
                if tele[3]: tele[3] = tele_effect(tele[3], tele_x[3], lucky3 * 3, 450)

                nv54 = nhanVat("nv54", w54_name[int(wi15)], w14_x, 550, 100, 100)
                nv54.draw_nv(4)
                if nv54.check_vc(box_but4) and box_touch[4] == 0: w14_x, x4 = lucky_atr(w14_x, x4, lucky4, nv54, 4)
                if k_stun[4]: w14_x, k_stun[4] = stun_effect(w14_x, x4, k_stun[4], 550)
                if rever[4]: w14_x -= 2 * x4
                if tele[4]: tele[4] = tele_effect(tele[4], tele_x[4], lucky4 * 3, 550)

                nv55 = nhanVat("nv55", w55_name[int(wi15)], w15_x, 650, 100, 100)
                nv55.draw_nv(5)
                if nv55.check_vc(box_but5) and box_touch[5] == 0: w15_x, x5 = lucky_atr(w15_x, x5, lucky5, nv55, 5)
                if k_stun[5]: w15_x, k_stun[5] = stun_effect(w15_x, x5, k_stun[5], 650)
                if rever[5]: w15_x -= 2 * x5
                if tele[5]: tele[5] = tele_effect(tele[5], tele_x[5], lucky5 * 3, 650)

                if check_press(exit_but.image_rect, pos):
                    off_screen_except(0.5)

            list_x = [w11_x, w12_x, w13_x, w14_x, w15_x]

            for i in range(5):
                if list_x[i] > finish + 30 and list_nv_win[i] == 0:
                    list_ranking.append("set05/PNG" + str(i + 1) + ".png")
                    list_stt.append(i + 1)
                    list_nv_win[i] = 1

            if len(list_stt):
                if setnv12:
                    sttnv = 10
                    map = 2
                    set = 1

                if setnv22:
                    sttnv = 15
                    map = 2
                    set = 2
                if nv2 == list_stt[0] and get_coin == False:
                    coin = int((int(coin) - int(bet)) + int(bet) * 4)
                    write_history(str(nv2 + sttnv), "win", bet, 3 * int(bet), "Main game")
                    win_or_lose = True
                    get_coin = True
                if nv2 != list_stt[0] and get_coin == False:
                    coin = int((coin - int(bet)))
                    write_history(str(nv2 + sttnv), "lose", bet, -int(bet), "Main game")
                    get_coin = True
                    win_or_lose = False

                if check_press(exit_but.image_rect, pos):
                    off_screen_except(0.5)
            if min(w11_x, w12_x, w13_x, w14_x, w15_x) > finish + 30:
                s = podium.after_race(list_stt[0], list_stt[1], list_stt[2], list_stt[3], list_stt[4], choosen_set,
                                      screen,
                                      SCREEN_WIDTH, SCREEN_HEIGHT, win_or_lose)
                if s == 22:
                    off_screen_except(22)

    if set1:
        current_time = pygame.time.get_ticks()
        elapsed_time = current_time - start_time
        if elapsed_time < 1000:
            background1.draw_bg(screen)
            pygame.mixer.music.stop()
            pygame.mixer.music.load("music/countdown1.mp3")
            pygame.mixer.music.play(0)
        elif elapsed_time < three_limit + 1000:
            background1.draw_bg(screen)
            three_but.draw_but(screen)
        elif elapsed_time < two_limit + 1000:
            background1.draw_bg(screen)
            two_but.draw_but(screen)
        elif elapsed_time < one_limit + 1000:
            background1.draw_bg(screen)
            one_but.draw_but(screen)
        elif elapsed_time < time_limit + 1000:
            background1.draw_bg(screen)
            if lan:
                start_but.draw_but(screen)
            else:
                start_butvn.draw_but(screen)
            pygame.mixer.music.stop()
            pygame.mixer.music.load("music/lucvodua2.mp3")
            pygame.mixer.music.play(0)
            w11_x = w12_x = w13_x = w14_x = w15_x = 1
        else:

            boom_but1 = Button(boom, box_but1.x, 225, 90, 90)
            boom_but2 = Button(boom, box_but2.x, 345, 90, 90)
            boom_but3 = Button(boom, box_but3.x, 470, 90, 90)
            boom_but4 = Button(boom, box_but4.x, 570, 90, 90)
            boom_but5 = Button(boom, box_but5.x, 660, 90, 90)
            map1(pos)
            if check_press(exit_but.image_rect, pos):
                pygame.mixer.music.stop()
                pygame.mixer.music.load("music/nhacnengame.mp3")
                pygame.mixer.music.play(-1)
                off_screen_except(0.5)
                time.sleep(0.2)

            if setnv21:
                choosen_set = 4
                nv41 = nhanVat("nv41", w41_name[int(wi14)], w11_x, 200, 120, 120)
                nv41.draw_nv(1)
                if nv41.check_vc(box_but1) and box_touch[1] == 0: w11_x, x1 = lucky_atr(w11_x, x1, lucky1, nv41, 1)
                if k_stun[1]: w11_x, k_stun[1] = stun_effect(w11_x, x1, k_stun[1], 200)
                if rever[1]: w11_x -= 2 * x1
                if tele[1]: tele[1] = tele_effect(tele[1], tele_x[1], lucky1 * 3, 200)

                nv42 = nhanVat("nv42", w42_name[int(wi14)], w12_x, 325, 120, 120)
                nv42.draw_nv(2)
                if nv42.check_vc(box_but2) and box_touch[2] == 0: w12_x, x2 = lucky_atr(w12_x, x2, lucky2, nv42, 2)
                if k_stun[2]: w12_x, k_stun[2] = stun_effect(w12_x, x2, k_stun[2], 325)
                if rever[2]: w12_x -= 2 * x2
                if tele[2]: tele[2] = tele_effect(tele[2], tele_x[2], lucky2 * 3, 325)

                nv43 = nhanVat("nv43", w43_name[int(wi14)], w13_x, 450, 120, 120)
                nv43.draw_nv(3)
                if nv43.check_vc(box_but3) and box_touch[3] == 0: w13_x, x3 = lucky_atr(w13_x, x3, lucky3, nv43, 3)
                if k_stun[3]: w13_x, k_stun[3] = stun_effect(w13_x, x3, k_stun[3], 450)
                if rever[3]: w13_x -= 2 * x3
                if tele[3]: tele[3] = tele_effect(tele[3], tele_x[3], lucky3 * 3, 450)

                nv44 = nhanVat("nv44", w44_name[int(wi14)], w14_x, 550, 120, 120)
                nv44.draw_nv(4)
                if nv44.check_vc(box_but4) and box_touch[4] == 0: w14_x, x4 = lucky_atr(w14_x, x4, lucky4, nv44, 4)
                if k_stun[4]: w14_x, k_stun[4] = stun_effect(w14_x, x4, k_stun[4], 550)
                if rever[4]: w14_x -= 2 * x4
                if tele[4]: tele[4] = tele_effect(tele[4], tele_x[4], lucky4 * 3, 550)

                nv45 = nhanVat("nv45", w45_name[int(wi14)], w15_x, 650, 120, 120)
                nv45.draw_nv(5)
                if nv45.check_vc(box_but5) and box_touch[5] == 0: w15_x, x5 = lucky_atr(w15_x, x5, lucky5, nv45, 5)
                if k_stun[5]: w15_x, k_stun[5] = stun_effect(w15_x, x5, k_stun[5], 650)
                if rever[5]: w15_x -= 2 * x5
                if tele[5]: tele[5] = tele_effect(tele[5], tele_x[5], lucky5 * 3, 650)

                list_x = [w11_x, w12_x, w13_x, w14_x, w15_x]

            if setnv11:
                choosen_set = 1
                map1(pos)
                nv11 = nhanVat("nv11", w11_name[int(wi11)], w11_x, 200, 120, 120)
                nv11.draw_nv(1)
                if nv11.check_vc(box_but1) and box_touch[1] == 0: w11_x, x1 = lucky_atr(w11_x, x1, lucky1, nv11, 1)
                if k_stun[1]: w11_x, k_stun[1] = stun_effect(w11_x, x1, k_stun[1], 200)
                if rever[1]: w11_x -= 2 * x1
                if tele[1]: tele[1] = tele_effect(tele[1], tele_x[1], lucky1 * 3, 200)

                nv12 = nhanVat("nv12", w12_name[int(wi11)], w12_x, 325, 120, 120)
                nv12.draw_nv(2)
                if nv12.check_vc(box_but2) and box_touch[2] == 0: w12_x, x2 = lucky_atr(w12_x, x2, lucky2, nv12, 2)
                if k_stun[2]: w12_x, k_stun[2] = stun_effect(w12_x, x2, k_stun[2], 325)
                if rever[2]: w12_x -= 2 * x2
                if tele[2]: tele[2] = tele_effect(tele[2], tele_x[2], lucky2 * 3, 325)

                nv13 = nhanVat("nv13", w13_name[int(wi11)], w13_x, 450, 120, 120)
                nv13.draw_nv(3)
                if nv13.check_vc(box_but3) and box_touch[3] == 0: w13_x, x3 = lucky_atr(w13_x, x3, lucky3, nv13, 3)
                if k_stun[3]: w13_x, k_stun[3] = stun_effect(w13_x, x3, k_stun[3], 450)
                if rever[3]: w13_x -= 2 * x3
                if tele[3]: tele[3] = tele_effect(tele[3], tele_x[3], lucky3 * 3, 450)

                nv14 = nhanVat("nv14", w14_name[int(wi11)], w14_x, 550, 120, 120)
                nv14.draw_nv(4)
                if nv14.check_vc(box_but4) and box_touch[4] == 0: w14_x, x4 = lucky_atr(w14_x, x4, lucky4, nv14, 4)
                if k_stun[4]: w14_x, k_stun[4] = stun_effect(w14_x, x4, k_stun[4], 550)
                if rever[4]: w14_x -= 2 * x4
                if tele[4]: tele[4] = tele_effect(tele[4], tele_x[4], lucky4 * 3, 550)

                nv15 = nhanVat("nv15", w15_name[int(wi11)], w15_x, 650, 120, 120)
                nv15.draw_nv(5)
                if nv15.check_vc(box_but5) and box_touch[5] == 0: w15_x, x5 = lucky_atr(w15_x, x5, lucky5, nv15, 5)
                if k_stun[5]: w15_x, k_stun[5] = stun_effect(w15_x, x5, k_stun[5], 650)
                if rever[5]: w15_x -= 2 * x5
                if tele[5]: tele[5] = tele_effect(tele[5], tele_x[5], lucky5 * 3, 650)

            list_x = [w11_x, w12_x, w13_x, w14_x, w15_x]

            for i in range(5):
                if list_x[i] > finish + 30 and list_nv_win[i] == 0:
                    list_ranking.append("set01/PNG" + str(i + 1) + ".png")
                    list_stt.append(i + 1)
                    list_nv_win[i] = 1

            if len(list_stt):
                if setnv11:
                    sttnv = 5
                    map = 1
                    set = 1
                if setnv21:
                    sttnv = 0
                    map = 1
                    set = 2
                if nv == list_stt[0] and get_coin == False:
                    coin = int((int(coin) - int(bet)) + int(bet) * 4)
                    write_history(str(int(nv + sttnv)), "win", bet, 3 * int(bet), "Main game")
                    get_coin = True
                    win_or_lose = True
                if nv != list_stt[0] and get_coin == False:
                    coin = int((coin - int(bet)))
                    write_history(str(int(nv + sttnv)), "lose", bet, -int(bet), "Main game")
                    get_coin = True
                    win_or_lose = False

            if min(w11_x, w12_x, w13_x, w14_x, w15_x) > finish + 30:
                s = podium.after_race(list_stt[0], list_stt[1], list_stt[2], list_stt[3], list_stt[4], choosen_set,
                                      screen,
                                      SCREEN_WIDTH, SCREEN_HEIGHT, win_or_lose)
                if s == 22:
                    off_screen_except(22)

    if set3:
        map3(pos)
        if check_press(exit_but.image_rect, pos):
            off_screen_except(0.5)
            time.sleep(0.2)
        current_time = pygame.time.get_ticks()
        elapsed_time = current_time - start_time
        if elapsed_time < 1000:
            background3.draw_bg(screen)
            pygame.mixer.music.stop()
            pygame.mixer.music.load("music/countdown1.mp3")
            pygame.mixer.music.play(0)
        elif elapsed_time < three_limit + 1000:
            background3.draw_bg(screen)
            three_but.draw_but(screen)
        elif elapsed_time < two_limit + 1000:
            background3.draw_bg(screen)
            two_but.draw_but(screen)
        elif elapsed_time < one_limit + 1000:
            background3.draw_bg(screen)
            one_but.draw_but(screen)
        elif elapsed_time < time_limit + 1000:
            background3.draw_bg(screen)
            if lan:
                start_but.draw_but(screen)
            else:
                start_butvn.draw_but(screen)
            pygame.mixer.music.stop()
            pygame.mixer.music.load("music/lucvodua2.mp3")
            pygame.mixer.music.play(0)
            w11_x = w12_x = w13_x = w14_x = w15_x = 1
        else:

            car1_obj = nhanVat("car1_obj", car1, w11_x, 225, 140, 90)
            car1_obj.draw_nv(1)
            rotate("setbg3/set03/banhxe1.png", w11_x - 33, 243, 49, 49)
            rotate("setbg3/set03/banhxe1.png", w11_x + 32, 243, 49, 49)
            if car1_obj.check_vc(box_but1) and box_touch[1] == 0: w11_x, x1 = lucky_atr(w11_x, x1, lucky1, car1_obj, 1)
            if k_stun[1]: w11_x, k_stun[1] = stun_effect(w11_x, x1, k_stun[1], 225)
            if rever[1]: w11_x -= 2 * x1
            if tele[1]: tele[1] = tele_effect(tele[1], tele_x[1], lucky1 * 3, 225)

            car2_obj = nhanVat("car2_obj", car2, w12_x, 325, 160, 90)
            car2_obj.draw_nv(2)
            rotate("setbg3/set03/banhxe1.png", w12_x - 23, 350, 25, 25)
            rotate("setbg3/set03/banhxe1.png", w12_x + 12, 350, 25, 25)
            if car2_obj.check_vc(box_but2) and box_touch[2] == 0: w12_x, x2 = lucky_atr(w12_x, x2, lucky2, car2_obj, 2)
            if k_stun[2]: w12_x, k_stun[2] = stun_effect(w12_x, x2, k_stun[2], 325)
            if rever[2]: w12_x -= 2 * x2
            if tele[2]: tele[2] = tele_effect(tele[2], tele_x[2], lucky2 * 3, 325)

            car3_obj = nhanVat("car3_obj", car3, w13_x, 450, 160, 90)
            car3_obj.draw_nv(3)
            rotate("setbg3/set03/banhxe1.png", w13_x - 23, 470, 40, 40)
            rotate("setbg3/set03/banhxe1.png", w13_x + 20, 470, 40, 40)
            if car3_obj.check_vc(box_but3) and box_touch[3] == 0: w13_x, x3 = lucky_atr(w13_x, x3, lucky3, car3_obj, 3)
            if k_stun[3]: w13_x, k_stun[3] = stun_effect(w13_x, x3, k_stun[3], 450)
            if rever[3]: w13_x -= 2 * x3
            if tele[3]: tele[3] = tele_effect(tele[3], tele_x[3], lucky3 * 3, 450)

            rotate("setbg3/set03/banhxe2.png", w14_x - 20, 571, 30, 30)
            rotate("setbg3/set03/banhxe2.png", w14_x + 20, 571, 30, 30)
            car4_obj = nhanVat("car4_obj", car4, w14_x, 550, 170, 100)
            car4_obj.draw_nv(4)

            if car4_obj.check_vc(box_but4) and box_touch[4] == 0: w14_x, x4 = lucky_atr(w14_x, x4, lucky4, car4_obj, 4)
            if k_stun[4]: w14_x, k_stun[4] = stun_effect(w14_x, x4, k_stun[4], 550)
            if rever[4]: w14_x -= 2 * x4
            if tele[4]: tele[4] = tele_effect(tele[4], tele_x[4], lucky4 * 3, 550)

            car5_obj = nhanVat("car5_obj", car5, w15_x, 675, 160, 90)
            car5_obj.draw_nv(5)
            rotate("setbg3/set03/banhxe1.png", w15_x - 42, 703, 40, 40)
            rotate("setbg3/set03/banhxe1.png", w15_x + 13, 703, 40, 40)
            rotate("setbg3/set03/banhxe1.png", w15_x + 60, 703, 40, 40)
            if car5_obj.check_vc(box_but5) and box_touch[5] == 0: w15_x, x5 = lucky_atr(w15_x, x5, lucky5, car5_obj, 5)
            if k_stun[5]: w15_x, k_stun[5] = stun_effect(w15_x, x5, k_stun[5], 650)
            if rever[5]: w15_x -= 2 * x5
            if tele[5]: tele[5] = tele_effect(tele[5], tele_x[5], lucky5 * 3, 650)

            list_x = [w11_x, w12_x, w13_x, w14_x, w15_x]

            for i in range(5):
                if list_x[i] > finish + 30 and list_nv_win[i] == 0:
                    '''list_ranking.append("set01/PNG" + str(i + 1) + ".png")'''
                    list_stt.append(i + 1)
                    list_nv_win[i] = 1

            if len(list_stt):
                if nv3 == list_stt[0] and get_coin == False:
                    coin = int((int(coin) - int(bet)) + int(bet) * 4)
                    write_history(str(int(nv3 + 20)), "win", bet, 3 * int(bet), "Main game")
                    get_coin = True
                    map = 3
                    win_or_lose = True
                if nv3 != list_stt[0] and get_coin == False:
                    coin = int((coin - int(bet)))
                    write_history(str(int(nv3 + 20)), "lose", bet, -int(bet), "Main game")
                    get_coin = True
                    map = 3
                    win_or_lose = False

            if min(w11_x, w12_x, w13_x, w14_x, w15_x) > finish + 30:
                s = podium.after_race(list_stt[0], list_stt[1], list_stt[2], list_stt[3], list_stt[4], 3,
                                      screen,
                                      SCREEN_WIDTH, SCREEN_HEIGHT, win_or_lose)
                off_screen_except(s)

    if lang:
        language_screen()
    if minigame_bool:
        minigame_screen()
    if rule_bool:
        Rule.rule_screen(lan, rule_page + 1)

        exit_but.draw_but(screen)
        if check_press(exit_but.image_rect, pos):
            off_screen_except(0)
        if pygame.mouse.get_pressed()[0] and rule_time:
            rule_page += 1
            rule_page %= 2
            rule_time = 0

        if pygame.mouse.get_pressed()[0] == 0 and rule_time == False:
            rule_time = 1
    if choosenv_bool1:
        choose_nv_screen1()
    if choosenv_bool2:
        choose_nv_screen2()
    if choosenv_bool3:
        choose_nv_screen3()

        if check_press(exit_but.image_rect, pos):
            time.sleep(0.2)
            off_screen_except(12)
        if check_press(pygame.Rect(50, 200, 200, 80), pos):
            time.sleep(0.2)
            off_screen_except(0.5)

    if ketqua:
        ketqua_screen(map, set, list_stt)

    if history_bool:

        if check_press(exit_but.image_rect, pos):
            off_screen_except(0)

        font = pygame.font.Font(None, 24)
        text_list = read_file2('account/' + str(user) + '/history.txt')
        scroll_y = max(0, min(scroll_y, len(text_list) * 20 * len(max(text_list, key=len))))
        draw_text_list(screen, text_list, font, scroll_y)
        exit_but.draw_but(screen)

    pygame.display.update()
    clock.tick(60)
    with open("account/" + str(user) + '/coin.txt', 'w') as c:
        c.write(str(coin))
